print('\nStarting User Manual Automation Process...\n\nLoading Libraries...')
import tkinter as tk
from tkinter import filedialog

import cv2, fitz, io, re, os, time,pyautogui,webbrowser,pyperclip
import numpy as np
import matplotlib.pyplot as plt

from PyPDF2 import PdfReader
from PIL import Image, ImageOps
from docx import Document

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import warnings
warnings.filterwarnings("ignore")

import logging

from paddleocr import PaddleOCR
logging.getLogger('ppocr').setLevel(logging.WARNING)

print('Done\n')

def get_image_blocks(image_path):
    # Load the image
    image = cv2.imread(image_path)

#     # **Zoom the entire image first**
    zoom_factor = 1.5  # Set your desired zoom factor
    zoomed_image = cv2.resize(image, None, fx=zoom_factor, fy=zoom_factor, interpolation=cv2.INTER_CUBIC)
    
#     # Convert to grayscale
    gray = cv2.cvtColor(zoomed_image, cv2.COLOR_BGR2GRAY)
    
    # Apply adaptive thresholding to isolate tables or other objects of interest
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2)
    
    # Optional: You can use Canny edge detection to improve contour detection
    # edges = cv2.Canny(gray, 50, 150)
    # Use 'edges' in place of 'thresh' if using Canny
    
    # Find contours using RETR_EXTERNAL or another mode based on your use case
    contours, _ = cv2.findContours(thresh, cv2.RETR_LIST, cv2.CHAIN_APPROX_SIMPLE)
    
    # List to store final contours that are not inside others
    final_contours = []
    
    # Loop through all contours
    for i, contour1 in enumerate(contours):
        keep_contour = True
        x1, y1, w1, h1 = cv2.boundingRect(contour1)  # Bounding box of contour1
        
        # Compare with every other contour
        for j, contour2 in enumerate(contours):
            if i != j:
                x2, y2, w2, h2 = cv2.boundingRect(contour2)  # Bounding box of contour2
                
                # Check if contour1 is completely inside contour2
                if (x1 >= x2 and y1 >= y2 and x1 + w1 <= x2 + w2 and y1 + h1 <= y2 + h2):
                    keep_contour = False
                    break  # No need to check further if it's inside another contour
        
        if keep_contour:
            final_contours.append(contour1)
        # List to hold detected columns
    columns = []
    new_contour = []
    # Loop through detected contours
    for contour in final_contours:
        # Get bounding box for each contour
        x, y, w, h = cv2.boundingRect(contour)
        
        # Filter out small contours based on width
        if w > 50:  # Adjust width threshold as needed
            # Extract column image and store
            new_contour.append(contour)
            column_img = zoomed_image[y:y+h, x:x+w]  # Use the zoomed image here
            
            # Optional: Apply additional preprocessing (e.g., sharpening)
            kernel = np.array([[0, -1, 0],
                               [-1, 5, -1],
                               [0, -1, 0]])  # Sharpening kernel
            sharpened = cv2.filter2D(column_img, -1, kernel)
    
            columns.append((x, sharpened))
    
    # Sort columns based on x-coordinate

    columns = columns[::-1]
    new_contour = new_contour[::-1]
    return columns, new_contour, zoomed_image

def get_callout_image(all_ocr, image_path, columns, new_contour, zoomed_image, character_threshold = 4):
    image_numbering = []
    for idx,result in enumerate(all_ocr):
        filtered_data = filter_list(result)
        filtered_data = [re.sub(r'[^A-Za-z ]', '', item) for item in filtered_data]
        filtered_data = list(set(filtered_data))
        if len(filtered_data) > character_threshold:
            image_numbering.append(idx)
            zoomed_image_copy = zoomed_image.copy()        
            x, y, w, h = cv2.boundingRect(new_contour[idx])
            cv2.rectangle(zoomed_image_copy, (x, y), (x + w, y + h), (0, 0, 255), 5)
            arrow_start = (x - 50, y - 50)  # Starting point of the arrow
            arrow_end = (x, y)  # Ending point of the arrow (top-left corner of the rectangle)
            cv2.arrowedLine(zoomed_image_copy, arrow_start, arrow_end, (0, 0, 255), 3, tipLength=0.2)  # Draw red arrow
            cv2.imwrite(image_path[:-4] + '_' + str(idx) + '.png', zoomed_image_copy)      
            
    return image_numbering

def get_callout_OCR_and_images(image_numbering, all_ocr_content, image_path, final_ocr, final_image_path):
    uploaded_image_num = 0
    image_ocr = []
    if len(image_numbering) == 0:
        image_ocr = all_ocr_content
        new_image_path = image_path
        final_ocr.append(image_ocr)
        final_image_path.append(new_image_path)
    else:
        for img_num in image_numbering:
            new_image_path = image_path[:-4] + "_" + str(img_num) + '.png'
            if len(image_numbering) == 1:
                image_ocr = all_ocr_content
            elif img_num == image_numbering[-1]:
                image_ocr = all_ocr_content[uploaded_image_num:]
            else:
                image_ocr = all_ocr_content[uploaded_image_num:img_num + 1]
                uploaded_image_num = img_num + 1
            final_ocr.append(image_ocr)
            final_image_path.append(new_image_path)
    return final_ocr, final_image_path


def generate_user_manual_by_ChatGPT(final_ocr):
#     chat_history = ['''New Question: I need a detailed user manual created from OCR data. Please transform the provided OCR text into a comprehensive and detailed user manual. When interpreting the OCR data, take care to understand the context, as the data might be sourced from graphs or other complex structures.''',
#                 '''Use '#' for heading level 1, '##' for heading level 2, and so forth. Do not use '####' or higher-level headers. Elaborate and explain the content in a detailed manner. Present the user manual topics as bullet points or numbered lists, avoiding any irrelevant or absurd points.''',
#                 '''I will provide the OCR text one page at a time. Please ensure each page is detailed and thoroughly explained. Do not focus on the values. Write the user manual so that any user can understand the content without needing the specific values. If you want to use values, include them only as examples. but use less examples.''',
#                 '''Do not write instructions, instead use it to write the user manual in detailed manner. and Dont mention the dates as well.'''
#             ]
    chat_history = ['''I need a detailed user manual created from OCR data. Please transform the provided OCR text into a comprehensive and detailed user manual. When interpreting the OCR data, take care to understand the context, as the data might be sourced from graphs or other complex structures. I am giving you the common instructions to follow for each page.1. Use '#' for heading level 1, '##' for heading level 2, and so forth. Do not use '####' or higher-level headers.2. Write don't in markdown format.3. Elaborate and explain the content in a detailed manner.4. Present the user manual topics as bullet points or numbered lists, avoiding any irrelevant or absurd points.5. I will provide the OCR text one page at a time. Please ensure each page is detailed and thoroughly explained.6. Do not focus on the values. Write the user manual so that any user can understand the content without needing the specific values. If you want to use values, include them only as examples.''']
    
     
    for idx, ocr  in enumerate(final_ocr):
        chat_history.append(f'''Please dont write it in Markdown. Here is the OCR of Page {idx + 1} of the user manual : {ocr}. Please ensure this page is detailed and thoroughly explained. Dont mention the dates and values as well. Use '#' for heading level 1, '##' for heading level 2, and so forth. Do not use '####' or higher-level headers and avoid any irrelevant or absurd points.''')
    
    chats = []
    # Open ChatGPT in the default web browser
    webbrowser.open('https://chat.openai.com/')
    time.sleep(10)  # Wait for the page to load (adjust if necessary)

    # Click on the chat window to focus it
      # Adjust the coordinates to click within the chat area

    screen_width, screen_height = pyautogui.size()
    screen_width = 1047
    screen_height = 454
    pyautogui.click(screen_width, screen_height)
    
    screen_width_copy = 420
    screen_height_copy = 652
    for num, input_text in enumerate(chat_history):
        pyautogui.click(screen_width, screen_height)

        # Type your input
        pyautogui.write(input_text, interval=0.01)
        pyautogui.press('enter')  # Press Enter to submit the input
        time.sleep(15)
        if num > 0:
            time.sleep(20)

            pyautogui.click(screen_width, screen_height)

            num_scrolls = 100

            for _ in range(num_scrolls):
                pyautogui.scroll(-100)  # Scroll up (positive value)
              # Click to ensure the chat window is focused

            # Use Tab key to navigate to the copy button
            
            pyautogui.click(screen_width_copy, screen_height_copy)

    #         pyautogui.press('tab', presses=1, interval=0.5)  # Adjust the number of presses as needed

    #         # Activate the button with Enter key
    #         pyautogui.press('enter')  # Press Enter to click the copy button

            # Wait briefly to ensure the clipboard has the new content
            time.sleep(1)
        
            # Use pyperclip to get the copied output
            output_text = pyperclip.paste()
            if (len(chats) > 0) and (output_text == chats[-1]):
                time.sleep(15)
                pyautogui.click(screen_width, screen_height)

                num_scrolls = 100

                for _ in range(num_scrolls):
                    pyautogui.scroll(-100)  # Scroll up (positive value)

                pyautogui.click(screen_width_copy, screen_height_copy)
                output_text = pyperclip.paste()
                
            if output_text.strip():  # Check if there's any text after stripping whitespace

                chats.append(output_text)
    pyautogui.hotkey('ctrl', 'w')
    return chats


# def generate_user_manual_by_ChatGPT(final_ocr):
#     chat_history = ['''I want to generate a detailed user manual documentation, i am giving you the OCR, please make a detailed user manual documentation for me.''',
#                 '''Use '#' for heading level 1, '##' for heading level 2, and so on, but Do not use '####', rather than use '**' and dont write it in markdown. and Can you please elaborate and explain the things in detailed manner.''',
#                 '''Please write the user manual topics in points. I will give the OCR of one page at a time.'''
#             ]
     
#     for idx, ocr  in enumerate(final_ocr):
#         chat_history.append(f'OCR of Page {idx + 1} of the user manual : {ocr}')
#         # chat_history.append('''Do not use '####' and dont write it in markdown. and Can you please elaborate and explain the things in detailed manner. Please write the user manual topics in points''')
    
#     chats = []
#     # Open ChatGPT in the default web browser
#     webbrowser.open('https://chat.openai.com/')
#     time.sleep(10)  # Wait for the page to load (adjust if necessary)

#     # Click on the chat window to focus it
#     screen_width, screen_height = pyautogui.size()
#     screen_width = screen_width *.371
#     screen_height = screen_height * .706
#     pyautogui.click(500, 500)
    
#     for num, input_text in enumerate(chat_history):


#         # Type your input
#         pyautogui.write(input_text, interval=0.01)
#         pyautogui.press('enter')  # Press Enter to submit the input
#         time.sleep(10)
#         if num > 2:
#             time.sleep(20)
            
#         # Wait for the response to be generated
#           # Get the copied text

#         # Check if the output text is empty
#         if (num >2):
#             pyautogui.click(500, 500)
#             num_scrolls = 100

#             for _ in range(num_scrolls):
#                 pyautogui.scroll(-100)  # Scroll up (positive value)
#               # Click to ensure the chat window is focused

#             # Use Tab key to navigate to the copy button
#             pyautogui.click(screen_width, screen_height)

#     #         pyautogui.press('tab', presses=1, interval=0.5)  # Adjust the number of presses as needed

#     #         # Activate the button with Enter key
#     #         pyautogui.press('enter')  # Press Enter to click the copy button

#             # Wait briefly to ensure the clipboard has the new content
#             time.sleep(1)

#             # Use pyperclip to get the copied output
#             output_text = pyperclip.paste()
#             if output_text.strip():  # Check if there's any text after stripping whitespace

#                 chats.append(output_text)
#     pyautogui.hotkey('ctrl', 'w')
#     return chats

def add_border_to_image(cell, border_color="B0B0B0"):
        # Set the border properties
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
    
        # Add border to the cell
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement('w:{}'.format(border_name))
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '6')  # Thickness of border
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)  # Border color (hex format)
            tcPr.append(border)

def add_heading(doc, text, level):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)

def add_paragraph(doc, text, bold=False):
    """Add a paragraph to the document with optional bold formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True

def format_text(text):
    """Format the input text and return structured content."""
    lines = text.strip().split('\n')
    formatted_content = []

    for line in lines:
        line = line.strip()  # Remove leading/trailing whitespace
        # Check for headings based on the number of leading '#' characters
        if '---' in line:
            continue
        elif line.startswith('# '):  # Heading level 1
            formatted_content.append(('heading', 1, line[2:]))  # Remove '# '
        elif line.startswith('## '):  # Heading level 2
            formatted_content.append(('heading', 2, line[3:]))  # Remove '## '
        elif line.startswith('### '):  # Heading level 3
            formatted_content.append(('heading', 3, line[4:]))  # Remove '### '
        
        # Check for headings that use the format "• **Function:**"
        elif (line.startswith('- **') or line.startswith('• **')) and (line.endswith(':**') or line.endswith('**:')):
            heading_text = line.replace('*','').replace('-','').strip()  # Remove '• ' and ':'
            formatted_content.append(('heading', 3, heading_text.replace('*','')))  # Treat as level 3 heading
            
        # Check for bold text
        elif line.endswith('**:') or line.endswith(':**'):  # Bold text
            formatted_content.append(('bold', line.replace('*','')))  # Remove '**'
        # Check for lines formatted like "4. **Upload Data File:**"
        elif line.count('.') == 1 and line.split('.')[0].isdigit() and line.strip().endswith(':'):
            # Split the line to get the heading text without the number
            heading_text = line.split('.')[1].strip()  # Get text after the number and remove extra spaces
            formatted_content.append(('heading', 3, heading_text.replace('*','')))  # Treat as level 3 heading
        # Check for bullet points
        elif line.startswith('* ') or line.startswith('- '):  # Bullet points
            formatted_content.append(('bullet', line[2:].replace('*','')))  # Remove '* ' or '- '
        # Check for numbered points
        elif len(line) > 1 and line[0].isdigit() and line[1] == '.':
            formatted_content.append(('numbered', line[3:].replace('*','')))  # Remove '1. '
        else:
            if line:  # Avoid adding empty paragraphs
                formatted_content.append(('paragraph', line.replace('*','')))

    return formatted_content

def create_document(formatted_content, doc):
    """Create a Word document with the provided formatted content."""
    

    for content in formatted_content:
        if content[0] == 'heading':
            add_heading(doc, content[2], content[1])
        elif content[0] == 'bold':
            add_paragraph(doc, content[1], bold=True)
        elif content[0] == 'bullet':
            # Add bullet points using ListBullet style
            para = doc.add_paragraph(style='ListBullet')  # Create a bullet point paragraph
            run = para.add_run(content[1])  # Add the bullet text
            run.bold = False  # Ensure bullet text is not bold
        elif content[0] == 'numbered':
            # Add numbered points using ListNumber style
            para = doc.add_paragraph(style='ListNumber')
            run = para.add_run(content[1])  # Add the numbered text
            run.bold = False  # Ensure numbered text is not bold
        elif content[0] == 'paragraph':
            add_paragraph(doc, content[1], bold=False)

    return doc



def formatting_texts(user_manual_doc, doc, pdf_path, generator, final_image_path, all_hashtags, page_width):
    
    for i,input_text in enumerate(user_manual_doc):
        formatted_content = format_text(input_text)

        # Create a new Word document
        doc = create_document(formatted_content, doc)
        
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(all_hashtags[i])
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(0.01)
        
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
    
        # Add the image to the cell and resize it
        image_path = final_image_path[i]
        
        
        run = cell.paragraphs[0].add_run()
        run.add_picture(image_path, width=page_width)
        
        # Add a border to the cell
        add_border_to_image(cell)

        # Save the document
    doc.save(str(pdf_path[:-4]) + '_' + generator + '.docx')
    print('User Manual Created. \nCheck ' + str(pdf_path[:-4]) + '_' + generator + '.docx')

def get_images(pdf_document, pdf_path):
    images = []
    zoom_x = 4.0  # Horizontal zoom
    zoom_y = 4.0  # Vertical zoom
    matrix = fitz.Matrix(zoom_x, zoom_y)  # For scaling
    bottom_margin = 100
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        image = page.get_pixmap(matrix=matrix)  # Apply the zoom factor
        
        # Convert the pixmap to an image object (Pillow)
        image_data = image.tobytes("png")
        image_pil = Image.open(io.BytesIO(image_data))
        width, height = image_pil.size
        image_pil = image_pil.crop((0, 0, width, height - bottom_margin))
        # Convert the image to grayscale
        grayscale_image = ImageOps.grayscale(image_pil)
        
        # Apply a threshold to binarize the image (turn into black and white)
        threshold = 240
        bw_image = grayscale_image.point(lambda x: 0 if x < threshold else 255, '1')
        
        # Find the last row with text
        width, height = bw_image.size
        pixels = bw_image.load()
        
        last_text_row = 0
        for y in range(height):
            if any(pixels[x, y] == 0 for x in range(width)):  # 0 means black pixel (text)
                last_text_row = y
        
        # Crop the image to include only up to the last row with text
        cropped_image = image_pil.crop((0, 0, width, last_text_row + 10))  # +10 for a little padding
    
        # Save the cropped image
        image_path = f'{pdf_path[:-4]}_img_{page_num + 1}.png'
        cropped_image.save(image_path)
        images.append(image_path)
    return images

def get_ocr_content(columns,  ocr):
    all_ocr = []
    for idx, (x, col) in enumerate(columns):
        new_ocr = []
        result = ocr.ocr(col, cls=True, det = True, rec = True)
        if result[0] != None:
            for i in result[0]:
                new_ocr.append(i[-1][0])
        all_ocr.append(new_ocr)
    
    return all_ocr

import re
def should_remove(item):

    date_pattern = r'\b(?:\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4}|\b\w{3,9}\s?\d{1,2}\b|\b\d{1,2}\s?\w{2}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b)\b'
    number_pattern = r'^\d+$'
    decimal_pattern = r'^\d{0,9}(,\d{0,9})*(\.\d+)?'
    negative_pattern = r'^-?\d{0,9}(,\d{0,9})*(\.\d+)?'
    dollar_pattern = r'\$'
    generic_k_pattern = r'\b[A-Za-z0-9]+[oO0-9]*[Kk]\b'
    
    return (re.fullmatch(date_pattern, item) or
            re.fullmatch(number_pattern, item) or
            re.fullmatch(decimal_pattern, item) or
            re.fullmatch(negative_pattern, item) or
            re.search(dollar_pattern, item) or
            re.fullmatch(generic_k_pattern, item) or
            re.fullmatch(generic_k_pattern, item[1:]))
    

# Filter function to apply on any list
def filter_list(input_list):
    return [item for item in input_list if not should_remove(item)]


def remove_subsets(lists, contour):
    result = []
    new_contour = []
    for idx,lst in enumerate(lists):
        if not any(set(lst).issubset(set(other)) for other in lists if lst != other):
            result.append(lst)
            new_contour.append(contour[idx])
    return result, new_contour

def generate_hashtags(filter_ocr):
    hashtags = []
    for idx, ocr_list in enumerate(filter_ocr):
        new_list = []
        for text in ocr_list:
            text = text.lower()
            text = re.sub(r'[^a-z]', '', text)
            text_len = len(text)
            if text_len > 5 and text_len < 25:
                new_list.append('#' + str(text))
        hashtags += new_list
    return hashtags

def get_final_ocr(final_ocr):
    hashtags = []
    for idx, ocr_list in enumerate(final_ocr):
        new_list = []
        for text in ocr_list:
            text = text.lower()
            text = re.sub(r'[^a-z]', ' ', text)
            text_len = len(text)
            text = str(text).strip()

            if text_len > 3 and text != '':
                new_list.append(text)
        hashtags += new_list
    return hashtags

def get_pdf_path():
    # Create a Tkinter root window (it will not be shown)
    root = tk.Tk()
    # Open the file dialog to select a PDF file
    file_path = filedialog.askopenfilename(
        title="Select a PDF file",
        filetypes=[("PDF files", "*.pdf")]  # Accept only PDF files
    )

    root.destroy()  # Close the root window after the dialog is closed

    return file_path

def create_user_manual():
    ocr = PaddleOCR(use_angle_cls=True, lang='en')
    page_width = Inches(5.9)
    doc = Document()
    pdf_path = get_pdf_path()
    character_threshold = 4

    print('Reading PDF Document...')
    pdf_document = fitz.open(pdf_path)
    print('Done\n')

    print('Creating Images...')
    images = get_images(pdf_document, pdf_path)
    print('Done\n')
    final_ocr = []
    final_image_path = []

    print('Cropping Images and Extracting Text From Images...')
    for image_path in images:
        columns, contour, zoomed_image = get_image_blocks(image_path)
        all_ocr = get_ocr_content(columns, ocr)
        all_ocr_content, new_contour = remove_subsets(all_ocr, contour)
        image_numbering = get_callout_image(all_ocr_content, image_path, columns, new_contour, zoomed_image, character_threshold)
        final_ocr, final_image_path = get_callout_OCR_and_images(image_numbering, all_ocr_content, image_path, final_ocr, final_image_path)

    print('Done\n')

    print('Cleaning Text...')
    all_hashtags = []
    for idx, idx_list in enumerate(final_ocr):
        final_ocr[idx] = [filter_list(data)  for data in final_ocr[idx]]
        final_ocr[idx] = get_final_ocr(final_ocr[idx])
        all_hashtags.append(generate_hashtags(idx_list))

    print('Done\n')
    print('This script will open ChatGPT to generate the user manual. \nPlease avoid interacting with your system during this process. \nThe ChatGPT window will close automatically once the task is complete.\n\n')
    time.sleep(7)

    
    user_manual_doc_chatgpt = generate_user_manual_by_ChatGPT(final_ocr)
    formatting_texts(user_manual_doc_chatgpt, doc, pdf_path, '_chatgpt', final_image_path, all_hashtags, page_width)

    folder_path = '/'.join(pdf_path.split('/')[:-1])
    for filename in os.listdir(folder_path):
        if filename.lower().endswith('.png'):
            os.remove(os.path.join(folder_path, filename))

def main():
    start = time.time()

    create_user_manual()
    end = time.time()
    print("Execution Time : ",int(end-start) , "s")
