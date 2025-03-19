from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Inches, RGBColor, Pt, Cm

from .BasicGenerator import BasicGenerator, PageSize
from .DocParagraph import DocParagraph
from .DocTable import DocTable
from .DocPicture import DocPicture
from .SeismicReportTemplate import SRTemplate
from .UtilFunctions import add_comma_in_num_str

class SeismicReport(BasicGenerator):
    def __init__(self):
        super().__init__()
        # 修改为A3图纸，横向，两栏
        self.change_paper_size(PageSize.A3_LANDSCAPE,2)
        # 修改纸张Margin，单位mm
        self.change_paper_margin(32,25,32,25)
        # 格式统一修改
        self.body_style.paragraph_format.line_spacing = Pt(22)
        
    def creat_doc(self):
        self.__add_info()
        self.__add_seismic_chapter()
        
        
    def __add_info(self):
        model_name = "TestModel"
        par_context = SRTemplate.FIRST_INFO(model_name)
        paragraph = DocParagraph(par_context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)
        
    def __add_seismic_chapter(self):
        chapter_index = 8
        sub_index = 1 
        self.__add_seismic_chapter_title(chapter_index)
        sub_index = self.__add_seismic_embedding(chapter_index,sub_index)
        sub_index = self.__add_project_mass(chapter_index, sub_index)
        sub_index = self.__add_period( chapter_index,sub_index)
        sub_index = self.__add_shear_mass_ratio( chapter_index,sub_index)
        sub_index = self.__add_shear_and_moment( chapter_index,sub_index)
        sub_index = self.__add_horizental_moment_ratio_for_column( chapter_index,sub_index)
        sub_index = self.__add_disp_and_drift( chapter_index,sub_index)
        sub_index = self.__add_horizental_stiffness_ratio( chapter_index,sub_index)
        sub_index = self.__add_rotation_ratio( chapter_index,sub_index)
        sub_index = self.__add_stiffness_mass_ratio( chapter_index,sub_index)
        sub_index = self.__add_shear_capacity_ratio( chapter_index,sub_index)
        sub_index = self.__add_wind_acc( chapter_index,sub_index)
        
        
        
    def __add_seismic_chapter_title(self,chapter_index :int):
        
        yjk_version = "6.0.0"
        
        current_context = SRTemplate.SEISMIC_CHAPTER_TITLE
        par_context = DocParagraph(current_context.title(chapter_index))
        par_context.par_level  = 1
        self.add_title(par_context,12,6)
        paragraph_texts = current_context.paragraph(chapter_index,yjk_version)
        for context in paragraph_texts[:-1]:
            paragraph = DocParagraph(context)
            paragraph.style = self.body_style
            self.add_paragraph(paragraph)

        text = paragraph_texts[-1]
        paragraph = DocParagraph(text)
        paragraph.style = self.body_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.first_line_indent = 0
        self.add_paragraph(paragraph)
        
        figure_title = current_context.picture(chapter_index)
        paragraph = DocParagraph(figure_title)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
   
    def __add_seismic_embedding(self, chapter_index:int, sub_index:int):
        
        current_context = SRTemplate.SEISMIC_EMBEDDING
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        context = current_context.paragraph(chapter_index,sub_index)[0]
        paragraph = DocParagraph(context)
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)
        
        table_title = current_context.table(chapter_index,sub_index)
        paragraph = DocParagraph(table_title)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
        
        table = DocTable(3,7)
        table.merge_cells(1,4,2,4)
        table.merge_cells(1,5,2,5)
        table.set_table_context(current_context.table_context) 
        self.add_table(table)
        
        
        return sub_index + 1
    
    def __add_project_mass(self, chapter_index:int, sub_index:int):
        
        total_mass = 125452
        total_area = 4345
        average_load = total_mass / total_area * 10
        
        current_context = SRTemplate.PROJECT_MASS
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        contexts = current_context.paragraph(
            chapter_index,
            sub_index,
            total_mass = add_comma_in_num_str(total_mass),
            total_area = add_comma_in_num_str(total_area),
            average_load = average_load
        )
        paragraph = DocParagraph(contexts[0])
        paragraph.style = self.body_style
        self.add_paragraph(paragraph)
        
        table_title = current_context.table(chapter_index,sub_index)
        paragraph = DocParagraph(table_title)
        paragraph.style = self.small_title_style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.add_paragraph(paragraph)
        
        table = DocTable(4,4)
        table.set_table_context(current_context.table_context(
            A = 1
        )) 
        self.add_table(table)
        
        
        return sub_index + 1
    
    def __add_period(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.PERIOD
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_shear_mass_ratio(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.SHEAR_MASS_RATIO
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_shear_and_moment(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.SHEAR_AND_MOMENT
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_horizental_moment_ratio_for_column(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.HORIZENTAL_MOMENT_RATIO_FOR_COLUMN
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_disp_and_drift(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.DISP_AND_DRIFT
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_horizental_stiffness_ratio(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.HORIZENTAL_STIFFNESS_RATIO
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_rotation_ratio(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.ROTATION_RATIO
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_stiffness_mass_ratio(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.STIFFNESS_MASS_RATIO
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_shear_capacity_ratio(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.SHEAR_CAPACITY_RATIO
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    def __add_wind_acc(self, chapter_index:int, sub_index:int):
        current_context = SRTemplate.WIND_ACC
        self.__insert_title_par_2(current_context,chapter_index,sub_index)
        
        
        return sub_index + 1
    
    
    def __insert_title_par_2(self, current_context, chapter_index, sub_index):
        par_context = DocParagraph(current_context.title(chapter_index,sub_index))
        par_context.par_level  = 2
        self.add_title(par_context,6 ,6 )
    
    
    
    
    
    