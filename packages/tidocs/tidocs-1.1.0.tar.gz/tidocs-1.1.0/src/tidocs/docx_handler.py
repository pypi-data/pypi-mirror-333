import io
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.shared import qn


def merge_word_docs_with_tables(
    main_doc_data: bytes,
    table_doc_data: bytes,
    marker_text: str = "TIDOCS_REPLACE_TABLE",
) -> bytes:
    """
    Merges tables from one Word document into another at specified marker locations,
    preserving hyperlinks, list formatting, and other document relationships.

    Args:
        main_doc_data (bytes): The main document binary data
        table_doc_data (bytes): The document containing tables binary data
        marker_text (str): The text to look for where tables should be inserted

    Returns:
        bytes: The merged document as binary data
    """
    # Load both documents from binary data
    main_doc = Document(io.BytesIO(main_doc_data))
    table_doc = Document(io.BytesIO(table_doc_data))

    # Create a mapping of relationship IDs between documents
    rel_map = {}

    # Copy hyperlink relationships from table_doc to main_doc
    for rel_id, rel in table_doc.part.rels.items():
        if (
            rel.reltype
            == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
        ):
            new_rel_id = main_doc.part.relate_to(
                rel._target, rel.reltype, rel.is_external
            )
            rel_map[rel_id] = new_rel_id

    main_numbering = main_doc.part.numbering_part.element
    table_numbering = table_doc.part.numbering_part.element

    # Create mapping for abstract numbering IDs
    abstract_num_map = {}

    # Copy abstract numbering definitions with new IDs
    for abstract_num in table_numbering.xpath(".//w:abstractNum"):
        old_abstract_id = abstract_num.get(qn("w:abstractNumId"))
        new_abstract_id = str(100000 + int(old_abstract_id))
        abstract_num.set(qn("w:abstractNumId"), new_abstract_id)
        abstract_num_map[old_abstract_id] = new_abstract_id
        main_numbering.append(abstract_num)

    # Copy and update concrete numbering instances
    for num in table_numbering.xpath(".//w:num"):
        old_num_id = num.get(qn("w:numId"))
        new_num_id = str(100000 + int(old_num_id))
        num.set(qn("w:numId"), new_num_id)

        # Update abstract numbering reference
        abstract_num_id = num.xpath(".//w:abstractNumId")[0]
        old_abstract_ref = abstract_num_id.get(qn("w:val"))
        if old_abstract_ref in abstract_num_map:
            abstract_num_id.set(qn("w:val"), abstract_num_map[old_abstract_ref])

        main_numbering.append(num)

    # Find all tables in the table document
    tables_to_insert = {}
    current_heading = None

    # Associate tables with their preceding headings
    for element in table_doc.element.body:
        if element.tag.endswith("p"):
            paragraph_text = element.text.strip()
            if paragraph_text:
                current_heading = paragraph_text
        elif element.tag.endswith("tbl"):
            if current_heading:
                # Deep copy the table element
                table_copy = parse_xml(element.xml)

                # Update relationship IDs in the copied table
                # Find all hyperlinks using the proper namespace approach
                for hyperlink in table_copy.xpath(".//w:hyperlink"):
                    old_rid = hyperlink.get(qn("r:id"))
                    if old_rid in rel_map:
                        hyperlink.set(qn("r:id"), rel_map[old_rid])

                # Preserve list formatting in the table
                for paragraph in table_copy.xpath(".//w:p"):
                    num_pr = paragraph.xpath(".//w:numPr")
                    if num_pr:
                        # Ensure numbering properties are preserved
                        for num in num_pr:
                            num_id = num.xpath(".//w:numId")
                            if num_id:
                                num_id[0].set(qn("w:val"), f"{num_id[0].val + 100000}")

                tables_to_insert[current_heading] = table_copy

    # Process the main document
    for paragraph in main_doc.paragraphs:
        if marker_text in paragraph.text:
            # Get the table associated with this marker
            if paragraph.text in tables_to_insert:
                # Insert table after the paragraph
                table_element = tables_to_insert[paragraph.text]
                paragraph._p.getparent().replace(paragraph._p, table_element)

    # Save the merged document to bytes
    output = io.BytesIO()
    main_doc.save(output)
    return output.getvalue()


def merge_documents(doc_data: bytes, table_data: bytes) -> bytes:
    """
    Merge two Word documents, inserting table_data into doc_data.

    Args:
        doc_data: Main document binary data
        table_data: Table document binary data

    Returns:
        Merged document binary data
    """
    try:
        merged_data = merge_word_docs_with_tables(doc_data, table_data)
        return merged_data
    except Exception as e:
        print(f"Error merging documents: {str(e)}")
        raise
