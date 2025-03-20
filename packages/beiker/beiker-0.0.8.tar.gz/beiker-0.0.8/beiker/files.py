import os,re,configparser,shutil,json,datetime,secrets
import pandas as pd
import fitz  # pip install PyMuPDF，用于处理PDF文件，通常以fitz作为别名引入
# from PySide6.QtWidgets import QMessageBox
from PIL import Image, ImageDraw
from.log import printLog
from openpyxl import load_workbook  
from openpyxl.utils.exceptions import IllegalCharacterError
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.utils import get_column_letter
from beiker.beiker_string import extract_file_archive_code

def find_missing_codes(files_df, records_df, files_df_column, records_df_column):
    """
    该函数用于检查 records_df 的 records_df_column 列经过 extract_file_archive_code 函数处理后的值，
    在 files_df 的 files_df_column 列中是否存在。如果不存在，则将处理后的值作为键，
    records_df 中对应单元格的地址列表作为值，存入字典并返回。

    参数:
    files_df (pandas.DataFrame): 包含文件信息的 DataFrame。
    records_df (pandas.DataFrame): 包含记录信息的 DataFrame。
    files_df_column (str): files_df 中用于比较的列名。
    records_df_column (str): records_df 中用于比较的列名。

    返回:
    dict: 键为经过 extract_file_archive_code 处理后在 files_df 中不存在的值，
          值为 records_df 中对应单元格的地址列表。
    """
    # 容错处理：检查输入的是否为有效的 DataFrame
    if not isinstance(files_df, pd.DataFrame) or not isinstance(records_df, pd.DataFrame):
        print("输入的参数不是有效的 DataFrame 类型，请检查输入。")
        return {}

    # 容错处理：检查列名是否存在于对应的 DataFrame 中
    if files_df_column not in files_df.columns or records_df_column not in records_df.columns:
        print("指定的列名不存在于对应的 DataFrame 中，请检查列名。")
        return {}

    # 获取 files_df 中指定列的唯一值集合
    existing_codes = set(files_df[files_df_column].apply(extract_file_archive_code))

    missing_codes_dict = {}

    # 遍历 records_df 的指定列
    for index, row in records_df.iterrows():
        record_value = row[records_df_column]
        processed_code = extract_file_archive_code(record_value)

        # 如果处理后的值不在 existing_codes 集合中
        if processed_code not in existing_codes:
            # 获取列的索引
            col_index = records_df.columns.get_loc(records_df_column) + 1
            # 生成 Excel 风格的单元格地址
            cell_address = f"{get_column_letter(col_index)}{index + 1}"

            if processed_code in missing_codes_dict:
                missing_codes_dict[processed_code].append(cell_address)
            else:
                missing_codes_dict[processed_code] = [cell_address]

    return missing_codes_dict

def tif_to_jpg(folder_path):
    """
    将指定文件夹中的所有 .tif 或 .tiff 文件转换为 .jpg 格式。
    
    参数:
        folder_path (str 或 list): 
            - 如果是 str，表示包含需要转换的 .tif 或 .tiff 文件的单个文件夹路径。
            - 如果是 list，表示包含多个文件夹路径的列表。
    """
    # 如果 folder_path 是一个列表，则遍历每个路径
    if isinstance(folder_path, list):
        paths = folder_path  # 使用列表中的路径
    else:
        paths = [folder_path]  # 将单个路径转换为列表

    # 遍历所有路径
    for path in paths:
        # 遍历指定文件夹中的所有文件
        for filename in os.listdir(path):
            # 检查文件扩展名是否是 .tif 或 .tiff（不区分大小写）
            if filename.lower().endswith(('.tif', '.tiff')):
                # 构造完整的文件路径
                tif_path = os.path.join(path, filename)
                # 打开.tif或.tiff文件
                with Image.open(tif_path) as img:
                    # 构造输出文件名，替换扩展名为.jpg
                    jpg_filename = os.path.splitext(filename)[0] + '.jpg'
                    jpg_path = os.path.join(path, jpg_filename)
                    # 保存为.jpg格式
                    img.convert('RGB').save(jpg_path, 'JPEG', quality=100)
                    print(f'Converted {tif_path} to {jpg_path}')

def process_date_accuracy(df):  
    """  
    处理 DataFrame 中的 date 字段，根据给定的模式为 date_accuracy 字段赋值。  
    
    参数:  
    df (pd.DataFrame): 输入的 DataFrame，必须包含 'date' 列。  

    返回:  
    pd.DataFrame: 处理后的 DataFrame，包含 'date_accuracy' 列。  
    """  
    # 定义正则表达式模式  
    pattern = r'^(?:[1-9]|1[0-2])月\d{1,2}日|' \
              r'^(?:[1-9]|1[0-2])月\d{1,2}日[至到](?:[1-9]|1[0-2])月\d{1,2}日|' \
              r'^(?:[1-9]|1[0-2])月\d{1,2}日[至到]\d{1,2}日$'  

    # 创建一个新列 'date_accuracy'，初始值设置为 2  
    df['date_accuracy'] = 2  

    # 遍历 DataFrame 的每一行  
    for index, row in df.iterrows():  
        date_value = row['date']  
        
        if pd.isna(date_value) or date_value == "":  
            # 如果 date 值为空，则赋值为 3  
            df.at[index, 'date_accuracy'] = 3  
        elif re.match(pattern, date_value):  
            # 如果 date 值符合模式，则赋值为 1  
            df.at[index, 'date_accuracy'] = 1  

    return df  

def validate_first_string(input_string):
    """
    函数功能：验证输入的字符串是否符合特定的日期或时间相关格式，并返回符合格式的第一个字符串部分（如果匹配成功），否则返回None。

    参数：
    input_string：需要验证的原始字符串。

    具体步骤如下：
    """
    # 将输入字符串中的英文逗号替换为中文逗号，以便后续按照统一的分隔符进行处理
    processed_string = input_string.replace(',', '，')
    # 以中文逗号为分隔符，将字符串分割成多个部分，得到一个字符串列表
    parts = processed_string.split('，')
    # 获取分割后的字符串列表中的第一个字符串，并去除其两端的空白字符（空格、制表符等）
    first_part = parts[0].strip()

    pattern = r'^(0?[1-9]|1[0-2])月|' \
              r'^(?:[1-9]|1[0-2])月\d{1,2}日|' \
              r'^(?:[1-9]|1[0-2])月\d{1,2}日[至到]\d{1,2}日|' \
              r'^(?:[1-9]|1[0-2])月\d{1,2}日[至到](?:[1-9]|1[0-2])月\d{1,2}日|' \
              r'^(?:[1-9]|1[0-2])月初|' \
              r'^\d{4}年(?:[1-9]|1[0-2])月[初中末]|' \
              r'^(?:[1-9]|1[0-2])月[上中下]旬|' \
              r'此后至\d{4}年|[春夏秋冬]季|' \
              r'暑假|寒假|暑期|([1-9]|1[0-2])月[至到]([1-9]|1[0-2])月|截至年底|' \
              r'“(一|二|三|四|五|六|七|八|九|十|[十一二]?[十]|[二]?[十一二]|三十|[三十][一|二|三|四|五|六|七|八|九]|[四][十][一|二|三|四|五|6]|[四][十][一|二|三|四|五|六|七|八|九])五”期间[\(（]\d{4}[～\-——~]\d{4}[\)）]|' \
              r'\d{1,2}月至\d{1,2}月|^\d{4}年[底初]?$'
    # 使用正则表达式对处理后的第一个字符串部分进行匹配
    if re.match(pattern, first_part):
        # 如果匹配成功，返回该符合格式的第一个字符串部分
        return first_part
    else:
        # 如果匹配失败，返回None
        return None

def process_dataframe(df, column_name):
    """
    函数功能：对给定的数据框（DataFrame）的指定列应用validate_first_string函数进行验证处理，
              然后将生成的'date'列移动到'year'列后面（如果'year'列存在），最后返回处理后的数据框。

    参数：
    df：需要处理的pandas DataFrame对象。
    column_name：DataFrame中要应用验证处理的列名，该列中的每个元素（字符串）会被传入validate_first_string函数进行处理。

    具体步骤如下：
    """
    # 对指定列（column_name）中的每个元素应用validate_first_string函数，
    # 将返回的结果（符合格式的字符串或None）组成一个新的Series，并赋值给名为'date'的新列
    df['date'] = df[column_name].apply(validate_first_string)

    # 将 'date' 列移动到 'year' 列后面（前提是 'year' 列在数据框的列名中存在）
    if 'year' in df.columns:
        # 获取当前数据框所有列的名称，并转换为列表形式，方便后续操作列顺序
        cols = df.columns.tolist()
        # 查找 'year' 列在列名列表中的索引位置
        year_index = cols.index('year')
        # 通过先弹出 'date' 列（pop操作会移除并返回该元素），再将其插入到 'year' 列后面的位置（year_index + 1），
        # 重新排列列的顺序，实现将 'date' 列移动到 'year' 列后面的目的
        cols.insert(year_index + 1, cols.pop())  # 将 'date' 放到 'year' 后面
        # 根据重新排列后的列顺序，重新构建数据框，此时 'date' 列就位于 'year' 列后面了
        df = df[cols]

    # 返回处理后的数据框（包含经过验证处理的 'date' 列以及调整好的列顺序，如果 'year' 列存在的话）
    return df

def split_text_between(original_df):  
    """  
    根据换行符分割original_df的text_between字段，将分割后的字符串和  
    对应的年份放入result_df中。  

    参数:  
    original_df (DataFrame): 包含'year'和'text_between'两列的原始DataFrame。  

    返回:  
    DataFrame: 包含'year'和'para'两列的结果DataFrame，  
                para列为分割后的字符串，year列为相应的年份。  
    """  
    
    # 初始化一个空列表来存储结果  
    results = []  

    # 遍历original_df中的每行  
    for index, row in original_df.iterrows():  
        year = row['year']  # 获取对应的年份  
        text_between = row['text_between']  # 获取对应的文本字符串  

        # 确保text_between为字符串类型，如果不是则跳过或处理为''（空字符串）  
        if isinstance(text_between, str):  
            # 按照换行符分割text_between字符串  
            paragraphs = text_between.split('\n')  

            # 为每个分割后的段落创建一个字典并追加到结果列表  
            for para in paragraphs:  
                results.append({  
                    'year': year,  # 将对应的年份放入结果中  
                    'para': para.strip()  # 去掉段落字符串两端的空白字符  
                })  
        else:  
            # 如果text_between不是字符串，可以选择跳过或记录日志  
            print(f"警告：第 {index} 行的 text_between 不是字符串，值为：{text_between}")  

    # 创建一个新的DataFrame 来存储结果  
    result_df = pd.DataFrame(results)  

    return result_df  

def extract_year_and_between_text(file_path):  
    """  
    在Word文件中查找类似"1952年"这样单独一行且居中的字符串，返回包含两列的DataFrame，  
    一列是找到的字符串（year），另一列是当前行year字段值之后直到下一行year字段值出现之间的文本（text_between）。  

    参数:  
    file_path (str): Word文件的路径  

    返回:  
    pd.DataFrame: 包含'year'和'text_between'两列的DataFrame  
    """  
    try:  
        doc = Document(file_path)  
        years = []  
        text_between = []  
        current_year = None  
        all_text_lines = []  
        first_year_found = False  # 标志，表示是否已经找到第一个年份  
        initial_text = []  # 存储第一个年份之前的文本  

        # 遍历文档段落，按行获取文本内容并查找符合条件的年份字符串  
        for para in doc.paragraphs:  
            lines = para.text.splitlines()  
            all_text_lines.extend(lines)  
            for line in lines:  
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:  
                    text = line.strip()  
                    if text.endswith("年") and text[:-1].isdigit():  
                        if not first_year_found:  
                            # 如果还没有找到年份，将之前的文本存储  
                            initial_text = "\n".join(all_text_lines[:-1])  # 取到当前行之前的所有文本  
                            first_year_found = True  
                        else:  
                            # 处理当前年份与上一年份之间的文本  
                            if current_year:  
                                start_index = None  
                                for i in range(len(all_text_lines) - 1, -1, -1):  
                                    if all_text_lines[i] == current_year:  
                                        start_index = i  
                                        break  
                                end_index = None  
                                for i in range(len(all_text_lines)):  
                                    if all_text_lines[i] == text:  
                                        end_index = i  
                                        break  
                                if start_index is not None and end_index is not None:  
                                    between_text_lines = all_text_lines[start_index + 1:end_index]  
                                    between_text = "\n".join(between_text_lines)  
                                    text_between.append(between_text)  
                        current_year = text  
                        years.append(current_year)  

        # 处理最后一个符合条件的年份字符串到文档末尾的文本  
        if current_year:  
            start_index = None  
            for i in range(len(all_text_lines) - 1, -1, -1):  
                if all_text_lines[i] == current_year:  
                    start_index = i  
                    break  
            if start_index is not None:  
                between_text_lines = all_text_lines[start_index + 1:]  
                between_text = "\n".join(between_text_lines)  
                text_between.append(between_text)  

        # 将初始文本添加到 DataFrame  
        years.insert(0, None)  # 在开头插入一个空值，表示没有年份  
        text_between.insert(0, initial_text)  # 将初始文本放在第一行  

        # 创建DataFrame并返回  
        result_df = pd.DataFrame({  
            "year": years,  
            "text_between": text_between  
        })  
        return result_df  
    except Exception as e:  
        print(f"出现错误: {e}")  
        return pd.DataFrame()  

def extract_text_around_keyword(text, keyword):
    """
    从给定字符串中查找指定关键词，并返回关键词前后字符串，直到遇到中文句号或者回车符为止。

    参数:
    text (str): 要进行查找操作的原始字符串
    keyword (str): 需要查找的关键词

    返回:
    list: 包含关键词及前后内容直到遇到中文句号或者回车符的字符串列表，如果未找到关键词则返回空列表
    """
    if isinstance(text, float):  # 判断text是否为float类型，如果是则转换为字符串类型
        text = str(text)
    if text == "":  # 明确判断text是否为空字符串
        return []
    result_texts = []
    start_index = 0
    while True:
        index = text.find(keyword, start_index)
        if index == -1:
            break
        # 向前查找中文句号或者回车符作为起始位置
        start = index
        while start > 0 and text[start - 1] not in ["。", "\n"]:
            start -= 1
        # 向后查找中文句号或者回车符作为结束位置
        end = index + len(keyword)
        while end < len(text) and text[end] not in ["。", "\n"]:
            end += 1
        result_texts.append(text[start:end + 1])
        start_index = index + 1
    return result_texts

def process_excel_and_extract_text(excel_file_path, keyword, output_folder_path):
    """
    读取指定Excel文件，查找每行数据中关键词相关文本，处理后保存为新的Excel文件。

    参数:
    excel_file_path (str): 要读取的Excel文件的路径
    keyword (str): 需要查找的关键词
    output_folder_path (str): 结果文件的输出文件夹路径

    返回:
    None
    """
    # 读取Excel文件到df中
    df = pd.read_excel(excel_file_path)
    result_df = pd.DataFrame()

    # 使用iterrows()方法遍历每一行进行处理
    for index, row in df.iterrows():
        rt = extract_text_around_keyword(row['text_between'], keyword)

        # 将rt列表转换为一个单列的DataFrame，列名设为'result'，同时添加'year'列
        rt_df = pd.DataFrame({'year': row['year'], 'result': rt})

        # 使用concat方法按列方向（axis=0）合并到原df中（按行合并，将新数据逐行添加到result_df中）
        result_df = pd.concat([result_df, rt_df], axis=0)

    # 根据关键词生成对应的输出文件名
    output_file_name = keyword + "result.xlsx"
    output_file_path = output_folder_path + '\\' + output_file_name if output_folder_path.endswith('\\') else output_folder_path + '\\' + output_file_name
    # 将结果保存到新的Excel文件中，不包含索引列
    result_df.to_excel(output_file_path, index=False)

def pdf_to_jpg(pdf_path, output_folder, dpi=300):
    """
    将指定的PDF文件的每一页转换为JPEG格式的图片，并保存到指定的输出文件夹中，可指定生成图片的分辨率（dpi）。

    参数:
    pdf_path (str): 要转换的PDF文件的路径。
    output_folder (str): 用于保存转换后JPEG图片的文件夹路径。
    dpi (int, 可选): 生成图片的分辨率，单位为每英寸点数（Dots Per Inch），默认值为300。

    返回:
    None
    """
    # 打开PDF文件
    pdf_document = fitz.open(pdf_path)
    # 遍历PDF的每一页
    for page_number in range(len(pdf_document)):
        # range(len(pdf_document))会根据PDF文件的总页数生成一个从0开始到总页数 - 1的整数序列，
        # 在每次循环中，page_number表示当前页面的序号（从0开始计数）。

        # 获取当前页面
        page = pdf_document[page_number]
        # 通过索引操作pdf_document[page_number]获取对应序号的页面对象，方便后续对每一页进行单独操作。

        # 将页面转换为图像
        pix = page.get_pixmap(dpi=dpi)
        # 调用页面对象page的get_pixmap()方法，传入dpi=dpi参数，按照设定的分辨率（默认300dpi或者用户指定的dpi值）
        # 将当前页面转换为一个图像数据对象pix，该对象包含了页面转换后图像的像素等相关信息，是后续生成和保存图片的基础。

        # 创建一个BytesIO对象
        img_bytes = io.BytesIO()
        # 创建一个BytesIO对象，它类似于一个内存中的文件对象，可以在内存中临时存储数据，
        # 这里用于暂存即将生成的JPEG格式图像的数据，方便后续进行保存操作，避免直接操作磁盘文件的频繁读写开销。

        # 将图像保存为JPEG格式
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        # 使用PIL库的frombytes()方法，从给定的像素数据创建一个Image对象。
        # "RGB"表示图像的色彩模式为红、绿、蓝三原色模式，[pix.width, pix.height]指定了图像的宽度和高度，
        # pix.samples是包含像素数据的字节序列，通过这些参数构建出一个可以进行保存操作的Image对象img。

        img.save(img_bytes, format='JPEG')
        # 调用Image对象img的save()方法，将图像数据保存到之前创建的BytesIO对象img_bytes中，
        # 并指定保存格式为'JPEG'，即将内存中的图像数据转换为JPEG格式存储在img_bytes中，方便后续将其写入磁盘文件。

        # 保存图像到指定文件夹
        img_bytes.seek(0)
        # 将BytesIO对象的内部指针移动到开头位置，因为之前保存图像数据到该对象后，指针处于末尾位置，
        # 而后续进行读取操作时需要从开头开始读取，所以要先将指针重置到起始位置。

        output_path = f"{output_folder}/page_{page_number + 1}.jpg"
        # 通过格式化字符串的方式生成图片的保存路径，文件名格式为page_页码.jpg，其中页码从1开始（page_number + 1），
        # 并与output_folder（输出文件夹路径）拼接起来，得到完整的图片保存路径output_path。

        with open(output_path, "wb") as img_file:
            img_file.write(img_bytes.read())
        # 使用Python的内置函数open()以二进制写入模式('wb')打开指定的output_path路径对应的文件，
        # 将img_bytes（存储着JPEG格式图像数据的内存对象）中的数据读取出来并写入到打开的文件中，
        # 从而实现将图像保存到指定文件夹下的磁盘文件中。

    # 关闭PDF文档
    pdf_document.close()

def read_excel_to_df(excel_file_path):
    """
    从指定的 Excel 文件路径读取数据到 DataFrame。

    参数：
    - excel_file_path：Excel 文件的路径。

    返回值：
    - df：如果读取成功，返回一个包含 Excel 数据的 DataFrame；如果出现错误，返回 None。
    """
    try:
        df = pd.read_excel(excel_file_path)
        return df
    except Exception as e:
        print(f"读取 Excel 文件时出现错误：{e}")
        return None
    
def process_excel(excel_path, remove_suffix):
    """
    该函数用于处理 Excel 文件，读取指定的 Excel 文件为 DataFrame，并删除以特定字符串结尾的列。

    参数：
    - excel_path：字符串类型，表示要处理的 Excel 文件路径。
    - remove_suffix：字符串类型，用于指定要删除的列名的结尾字符串。

    返回值：
    - processed_df：经过处理后的 DataFrame。

    步骤：
    1. 使用`pd.read_excel`读取指定路径的 Excel 文件，将其存储为 DataFrame 对象`df`。
    2. 使用列表推导式遍历`df`的列名，筛选出以`remove_suffix`结尾的列名，存储在`columns_to_drop`列表中。
    3. 使用`df.drop`方法删除`columns_to_drop`列表中的列，得到处理后的 DataFrame`processed_df`。
    4. 返回处理后的 DataFrame。
    """
    df = pd.read_excel(excel_path)
    columns_to_drop = [col for col in df.columns if col.endswith(remove_suffix)]
    processed_df = df.drop(columns=columns_to_drop)
    return processed_df
    
def merge_excels(excel_path_list, merge_excel_path, columns='all',qmessage=True):
    """
    该函数用于合并多个 Excel 文件为一个 Excel 文件。

    参数：
    - excel_path_list：列表类型，包含要合并的 Excel 文件路径列表。
    - merge_excel_path：字符串类型，表示合并后的 Excel 文件存储路径。
    - columns：字符串或列表类型。默认为'all'表示合并所有列；如果是列名的列表，则合并指定的列。

    返回值：
    - merge_excel_path：合并后的 Excel 文件存储路径。

    步骤：
    1. 创建一个空列表`all_data`，用于存储各个 Excel 文件读取后的 DataFrame。
    2. 遍历`excel_path_list`中的每个路径，使用`pd.read_excel`读取 Excel 文件。
        - 如果 columns 为'all'，直接将读取的 DataFrame 添加到`all_data`列表中。
        - 如果 columns 是一个列名的列表，选择指定的列添加到`all_data`列表中。
    3. 使用`pd.concat`将`all_data`列表中的所有 DataFrame 进行合并，设置`ignore_index=True`以重新生成连续的索引。
    4. 使用`merged_df.to_excel`将合并后的 DataFrame 保存到指定路径`merge_excel_path`，设置`index=False`表示不保存索引列。
    5. 返回合并后的 Excel 文件存储路径。
    """
    all_data = []
    for path in excel_path_list:
        if not os.path.exists(path):
            if qmessage:
                # QMessageBox.warning(None, '错误', f'文件 {path} 不存在。')
                print( f'merge_excels读取文件 {path} 不存在。')
            continue
        try:
            df = pd.read_excel(path)
            if columns == 'all':
                all_data.append(df)
            else:
                if isinstance(columns, list):
                    selected_df = df[columns]
                    all_data.append(selected_df)
                else:
                    raise ValueError("columns 参数必须是'all'或者列名的列表。")
        except Exception as e:
            print(f"读取文件 {path} 时出现错误：{e}")
    try:
        merged_df = pd.concat(all_data, ignore_index=True)
        merged_df.to_excel(merge_excel_path, index=False)
    except Exception as e:
        print(f"保存合并后的文件到 {merge_excel_path} 时出现错误：{e}")
    return merge_excel_path

def text_to_docx(filepath,text,head=''):
    # 创建一个新的 Word 文档  
    doc = Document()  
    # 添加标题  
    doc.add_heading(head, level=1)  
    # 添加字符串到文档  
    doc.add_paragraph(text)  
    # 保存文档  
    doc.save(filepath)  

def check_file_exists(file_path):
    """
    检查指定路径的文件是否存在。

    :param file_path: 文件的完整路径
    :return: 如果文件存在返回True，否则返回False
    """
    return os.path.exists(file_path)

def check_directory_exists(directory_path):
    """
    检查指定路径的目录是否存在。

    :param directory_path: 目录的完整路径
    :return: 如果目录存在返回True，否则返回False
    """
    return os.path.exists(directory_path) and os.path.isdir(directory_path)

def save_df_with_links(df, path_columns, excel_path):
    '''
    对存有df列名的path_columns列表中的列转化为链接，可以点击打开，并应用样式让数据更加具有可读性。
    '''
    print(f"开始处理df,并将数据存入: {excel_path}......")
    # 检查DataFrame是否为空
    if df.empty:
        print("save_df_with_links - DataFrame is empty. No need to save.")
        # QMessageBox.information(None, "信息", f"DataFrame is empty. No need to save.")
        return

    # 检查路径是否有效
    if not excel_path.endswith('.xlsx'):
        print("Invalid file path. Please provide a .xlsx file path.")
        return

    # 将DataFrame写入Excel文件
    try:
        df.to_excel(excel_path, index=False)
    except Exception as e:
        print(f"An error occurred while writing to Excel: {e}")
        # QMessageBox.critical(None, "错误",f"An error occurred while writing to Excel: {e}")
        return

    count = 0  # 对插入列进行计数
    try:
        # 加载Excel文件
        wb = load_workbook(excel_path)
        ws = wb.active
    except Exception as e:
        print(f"An error occurred while loading the Excel file: {e}")
        return

    # 忽略标题第一行，处理每个路径列
    for col in path_columns:
        if col in df.columns:
            # 获取当前列在Excel中的列索引（从1开始）
            col_index = df.columns.get_loc(col) + count + 1  # 获取列的索引并转为1-based
            
            # 在路径列后插入一个新列
            new_col_index = col_index + 1
            try:
                ws.insert_cols(new_col_index)
                ws.cell(row=1, column=new_col_index, value=f'{col}_Link')  # 新列标题
            except Exception as e:
                print(f"An error occurred while inserting columns: {e}")
                continue
            
            for row in range(2, ws.max_row + 1):  # 从第二行开始，因为第一行为标题
                path_value = ws.cell(row=row, column=col_index).value
                if path_value:
                    try:
                        # 创建超链接
                        open_file_cell = ws.cell(row=row, column=new_col_index)
                        open_file_cell.hyperlink = path_value
                        open_file_cell.value = 'Open File'
                        open_file_cell.font = Font(bold=True, color='FF0000')  # 红色加粗字体
                    except ValueError as e:
                        print(f"An error occurred while creating hyperlink: {e}")
                        continue
            count += 1
            
    # 设置标题行的样式
    title_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
    for col in ws.iter_cols(min_row=1, max_row=1, min_col=1, max_col=len(df.columns) + len(path_columns)):
        for cell in col:
            cell.fill = title_fill
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            
    # 设置每隔一行的背景颜色
    odd_row_fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
    for row in range(2, ws.max_row + 1):  # 从第二行开始，因为第一行为标题
        if row % 2 == 0:  # 偶数行（从2开始计数）
            for col in range(1, len(df.columns) + count + 1):
                ws.cell(row, col).fill = odd_row_fill

    # 冻结标题行
    if len(df.columns) > 0:
        ws.freeze_panes = "A2"

    # 保存修改后的Excel文件
    try:
        wb.save(excel_path)
        wb.close()
        print(f"df has been save to the Excel file: {excel_path}")
    except Exception as e:
        print(f"An error occurred while saving the Excel file: {e}")

def find_ori_files_from_json(file_path):
    '''
    将 other\0001.json 的 other 目录的上一级目录中，json 文件来源的文件找出来。
    found_files[0] 是路径，found_files[1] 是文件名
    '''
    try:
        # 获取文件所在的目录
        other_directory = os.path.dirname(file_path)
        # 获取 other 目录的上一级目录
        archive_directory = os.path.dirname(other_directory)
        # 获取文件名（不包括扩展名）
        file_name_without_extension = os.path.splitext(os.path.basename(file_path))[0]
        pattern = re.compile(r'(.*?)(?=-paddle|-ali)')
        result = [re.match(pattern, name).group(1) for name in [file_name_without_extension] if re.match(pattern, name)]

        if result:
            file_name_without_extension = result[0]
        else:
            raise ValueError("find_ori_files_from_json-Could not extract base filename.")

        # 初始化一个列表来存储找到的文件
        found_files = []
        # 遍历上一级目录中的所有文件
        for filename in os.listdir(archive_directory):

            # 检查文件名（不包括扩展名）是否与给定的文件名相同
            if os.path.splitext(filename)[0] == file_name_without_extension:
                # 构建完整的文件路径
                full_path = os.path.join(archive_directory, filename)

                # 将文件路径和文件名添加到列表中
                found_files.append(full_path)  # 只添加路径
                found_files.append(filename)  # 只添加文件名

        # 返回找到的文件
        return found_files
    except Exception as e:
        print(f"find_ori_files_from_json-An error occurred: {e}")
        return None

def get_parent_dir(file_path):
    """
    获取文件路径的上一层父目录路径。
    """
    # 获取上层文件夹路径
    parent_dir = os.path.dirname(file_path)
    return parent_dir

def read_json_file(file_path):
    try:
        print(f"read_json_file 正在处理文件： {file_path} ")
        with open(file_path, 'r', encoding='utf-8') as file:
            return json.load(file)
    except FileNotFoundError:
        print(f"read_json_file 文件 {file_path} 未找到。")
    except json.JSONDecodeError:
        print(f"read_json_file 文件 {file_path} 不是有效的JSON格式。")
    except Exception as e:
        print(f"read_json_file 读取文件 {file_path} 时发生错误: {e}")
    return None

def process_serial_numbers(serial_numbers_positions, extra_param, split_num):
    """
    此函数用于处理一系列的序列号位置信息，将其按照指定数量分组构建成DataFrame，并且用'NA'填充空值，同时添加额外的参数到每行。

    :param serial_numbers_positions: 包含序列号位置信息的列表，其中每个元素可能是一个包含序列号和其他相关信息的元组或列表，这里主要使用其中的第一个元素作为数据来源。
    :param extra_param: 额外的参数，将被添加到DataFrame的每一行中，可以是任何类型，只要能与其他数据一起构成DataFrame的行。
    :param split_num: 表示每多少个序列号位置信息组成DataFrame的一行，整数类型。

    函数内部逻辑如下：
    1. 首先创建一个空列表 `data`，用于存储构建DataFrame的数据。
    2. 外层循环通过 `range` 函数以步长为 `split_num` 遍历 `serial_numbers_positions` 列表的索引。
       - 对于每个循环，创建一个空列表 `row_data`，用于存储当前行的数据。
       - 内层循环从当前外层循环的索引 `i` 开始，到 `i + split_num`，遍历 `serial_numbers_positions` 列表。
         - 如果内层循环的索引 `j` 小于 `serial_numbers_positions` 列表的长度：
           - 如果当前位置的 `serial_numbers_positions` 元素长度大于0，则将其第一个元素添加到 `row_data` 中。
           - 如果当前位置的 `serial_numbers_positions` 元素为空，则添加'NA'到 `row_data` 中。
         - 如果内层循环的索引 `j` 大于等于 `serial_numbers_positions` 列表的长度，则跳出内层循环。
       - 在每行数据构建完成后，将额外参数 `extra_param` 添加到 `row_data` 中。
       - 如果当前行数据的长度等于 `split_num + 1`（考虑到额外参数占一个位置），则将该行数据添加到 `data` 列表中。
    3. 最后使用构建好的 `data` 列表创建一个DataFrame，列名由 `f'Column{i+1}' for i in range(split_num)` 和 `'ExtraParam'` 组成，然后返回这个DataFrame。
    """
    data = []
    for i in range(0, len(serial_numbers_positions), split_num):
        row_data = []
        for j in range(i, i + split_num):
            if j < len(serial_numbers_positions):
                if len(serial_numbers_positions[j]) > 0:
                    row_data.append(serial_numbers_positions[j][0])
                else:
                    row_data.append('NA')
            else:
                break
        row_data.append(extra_param)
        if len(row_data) == split_num + 1:
            data.append(row_data)
    df = pd.DataFrame(data, columns=[f'Column{i+1}' for i in range(split_num)] + ['ExtraParam'])
    return df


def check_file_in_folder(folder_path, file_name):
    """
    此函数用于检查指定文件夹中是否包含指定的文件。

    :param folder_path: 文件夹的路径，字符串类型，表示要检查的文件夹的位置。
    :param file_name: 要检查的文件名称，字符串类型，表示期望在文件夹中找到的文件名字。

    函数内部逻辑如下：
    1. 使用 `os.listdir` 函数获取指定文件夹中的所有文件列表，得到一个包含文件名的列表。
    2. 检查指定的文件名称是否在获取到的文件列表中。
       - 如果在列表中，则返回True，表示文件夹中包含指定文件。
       - 如果不在列表中，则返回False，表示文件夹中不包含指定文件。
    """
    files_in_folder = os.listdir(folder_path)
    if file_name in files_in_folder:
        return True
    else:
        return False

# 统计文件夹及子文件夹下的总文件数
# folder_paths = [r"C:\Users\picture2excel"]
# exclude_folder_regex_list = [r'^\d{4}年彩色$'] 按正则表达式过滤文件夹，要输入的是re.compile正则表达式对象
# exclude_extensions = ['.json', '.xlsx'] 均为小写即可
# exclude_folders = ['参考','alitest']
# only_extensions=[] ，进统计带此列表扩展名的文件
# file_only_suffix=[] ，仅统计文件名带给定后缀的文件，文件名指的是不带扩展名的文件名
# include_dir=False 统计包括文件夹
# first_dir=False 只统计当前第一级文件夹
# 返回也是列表，给定绝对/相对路径，就返回绝对/相对路径
def filtered_files(folder_paths,   
                   exclude_folders=[],   
                   exclude_extensions=[],   
                   exclude_files=[],   
                   exclude_folder_regex_list=[],   
                   only_file_regex_list=[],   
                   only_folder_regex_list=[],   #所有正则筛选只需输入正则表达式，不需要re.compile对象
                   exclude_file_regex_list=[],   
                   only_extensions=[],   
                   file_only_suffix=[],   #只统计具有指定结尾字符串的文件
                   folder_only_suffix=[],  # 只统计具有指定结尾字符串的的文件夹  
                   include_dir=False,   
                   first_dir=False):  
    # 将Regex字符串转换为相应的re.compile对象  
    exclude_folder_regex_list = [re.compile(regex) for regex in exclude_folder_regex_list]  
    only_folder_regex_list = [re.compile(regex) for regex in only_folder_regex_list]  
    exclude_file_regex_list = [re.compile(regex) for regex in exclude_file_regex_list]  
    only_file_regex_list = [re.compile(regex) for regex in only_file_regex_list]  

    # 用于统计总文件数  
    total_files = 0  
    # 用于存储文件名的列表  
    files_names = []  
    # 用于存储文件路径的列表  
    files_paths = []  

    # 遍历给定的文件夹路径列表  
    for folder_path in folder_paths:  
        # 如果 first_dir 为 True，只处理第一级  
        if first_dir:  
            root_files = []  
            root_names = []  
            root_paths = []  
            for item in os.listdir(folder_path):  
                item_path = os.path.join(folder_path, item)  
                if os.path.isdir(item_path):  
                    if include_dir:  
                        # 只统计符合 only_folder_regex_list 的文件夹  
                        if only_folder_regex_list and not any(regex.match(item) for regex in only_folder_regex_list):  
                            continue  
                        # 检查 folder_only_suffix  
                        if folder_only_suffix and not any(item.endswith(suffix) for suffix in folder_only_suffix):  
                            continue  
                        total_files += 1  
                        root_names.append(item)  
                        root_paths.append(item_path)  
                elif os.path.isfile(item_path):  
                    file_ext = os.path.splitext(item)[1].lower()  
                    # 使用 exclude_file_regex_list 来排除符合正则表达式的文件  
                    if exclude_file_regex_list and any(regex.match(item) for regex in exclude_file_regex_list):  
                        continue  
                    # 只统计符合 only_file_regex_list 的文件  
                    if only_file_regex_list and not any(regex.match(item) for regex in only_file_regex_list):  
                        continue  
                    if only_extensions and file_ext not in [ext.lower() for ext in only_extensions]:  
                        continue  
                    if any(file_ext == ext.lower() for ext in exclude_extensions):  
                        continue  
                    if item in exclude_files:  
                        continue  
                    if file_only_suffix and any(os.path.splitext(item)[0].endswith(suffix) for suffix in file_only_suffix):  
                        total_files += 1  
                        root_names.append(item)  
                        root_paths.append(item_path)  
                    elif not file_only_suffix:  
                        total_files += 1  
                        root_names.append(item)  
                        root_paths.append(item_path)  
            files_names = root_names  
            files_paths = root_paths  
        else:  
            for root, dirs, files in os.walk(folder_path):  
                for folder in exclude_folders:  
                    if folder in dirs:  
                        dirs.remove(folder)  
                if exclude_folder_regex_list:  
                    dirs[:] = [d for d in dirs if not any(regex.match(d) for regex in exclude_folder_regex_list)]  
                for folder in dirs:  
                    # 只统计符合 only_folder_regex_list 的文件夹  
                    if only_folder_regex_list and not any(regex.match(folder) for regex in only_folder_regex_list):  
                        dirs.remove(folder)  
                    # 检查 folder_only_suffix  
                    if folder_only_suffix and not any(folder.endswith(suffix) for suffix in folder_only_suffix):  
                        dirs.remove(folder)  
                if include_dir and root != folder_path:  
                    total_files += 1  
                    files_names.append(os.path.basename(root))  
                    files_paths.append(root)  
                for file in files:  
                    file_ext = os.path.splitext(file)[1].lower()  
                    # 使用 exclude_file_regex_list 来排除符合正则表达式的文件  
                    if exclude_file_regex_list and any(regex.match(file) for regex in exclude_file_regex_list):  
                        continue  
                    # 只统计符合 only_file_regex_list 的文件  
                    if only_file_regex_list and not any(regex.match(file) for regex in only_file_regex_list):  
                        continue  
                    if only_extensions and file_ext not in [ext.lower() for ext in only_extensions]:  
                        continue  
                    if any(file_ext == ext.lower() for ext in exclude_extensions):  
                        continue  
                    if file in exclude_files:  
                        continue  
                    if file_only_suffix and any(os.path.splitext(file)[0].endswith(suffix) for suffix in file_only_suffix):  
                        total_files += 1  
                        files_names.append(file)  
                        files_paths.append(os.path.join(root, file))  
                    elif not file_only_suffix:  
                        total_files += 1  
                        files_names.append(file)  
                        files_paths.append(os.path.join(root, file))  

    return total_files, files_names, files_paths

def split_file_name_extension_list(file_name_list):
    '''
    将文件全名列表，分割成文件名和不带点的扩展名两个列表，如果是文件夹，扩展名位置是空字符串。
    '''
    file_names = []
    extensions = []
    for name in file_name_list:
        parts = name.split('.')
        if len(parts) == 1:
            file_names.append(parts[0])
            extensions.append('')
        else:
            file_names.append('.'.join(parts[:-1]))
            extensions.append(parts[-1])
    extensions_with_dot = ['.' + ext for ext in extensions]
    return file_names, extensions_with_dot

# 函数功能：在data_value里基于正则表达式查找内容，处理结果存入DataFrame并添加file_path列
# data_value: 要进行正则查找的数据内容，数据类型可能是多种类型，只要能被json.dumps转换为JSON字符串即可，如字典、列表等
# file_path: 文件路径，字符串类型，将被用于后续操作
# columns: 要设置给DataFrame的列名，数据类型为可迭代对象，如列表，其中元素为字符串类型
# pattern: 用于查找的正则表达式模式，字符串类型
# split_num: 分割相关的参数（具体作用由process_serial_numbers函数决定），数据类型由process_serial_numbers函数内部的逻辑决定，可能是整数等类型
def regular_find_to_df(data_value, file_path, columns, pattern, split_num):
    # 在将data_value转换为JSON字符串（不进行ASCII编码转换）后，使用re.finditer在其中查找符合pattern的内容
    # re.finditer返回一个迭代器，每个元素包含匹配到的内容（match.group(0)）和匹配开始的位置（match.start()）
    matches = re.finditer(pattern, json.dumps(data_value, ensure_ascii=False))

    serial_numbers_positions = []

    # 遍历re.finditer返回的每个匹配结果
    for match in matches:
        # 将每个匹配结果中的匹配内容和开始位置作为元组添加到serial_numbers_positions列表中
        serial_numbers_positions.append((match.group(0), match.start()))

    # 取serial_numbers_positions列表中索引为23及之后的元素
    serial_numbers_positions = serial_numbers_positions[23:]

    # 使用process_serial_numbers函数处理serial_numbers_positions列表，同时传入文件路径相关信息和split_num参数
    # process_serial_numbers函数可能是自定义的，用于进一步处理匹配结果，这里得到一个DataFrame
    df = process_serial_numbers(serial_numbers_positions, split_path(os.path.abspath(file_path), '电子档案'), split_num)
    # 将传入的columns设置为DataFrame的列名
    df.columns = columns

    # 返回处理后的DataFrame
    return df

def split_path(path, folder_name):
    """
    将绝对路径转换为相对路径，不包含给定的文件夹名称。

    :param path: 要处理的绝对路径，数据类型为字符串，例如'/home/user/documents/folder1/file.txt'。
    :param folder_name: 用作参考的文件夹名称，数据类型为字符串，例如'folder1'。
    :return: 相对路径，数据类型为字符串，如果原路径中不存在folder_name则返回原路径。
    """
    if '/' in path:
        parts = path.split('/')
    else:
        parts = path.split('\\')
    try:
        index = parts.index(folder_name)
        relative_path_parts = parts[index + 1:]
        if '/' in path:
            return '/'.join(relative_path_parts)
        else:
            return '\\'.join(relative_path_parts)
    except ValueError:
        return path
    
def jsonAxes_to_excel(excel_file_name, org_image_path, prism_tables_info, columns_num):
    """
    从json文件中的prism_tables_info列表中提取单元格坐标和行数信息，并存入excel表。

    :param excel_file_name: 要保存的Excel文件的名称，数据类型为字符串，例如 'output.xlsx'。
    :param org_image_path: 原始图像的路径，数据类型为字符串，例如 '/path/to/image.jpg'，将被添加到Excel表中。
    :param prism_tables_info: 从json文件中获取的包含表格信息的列表，其中包含了表格的各种信息，例如每个单元格的信息等。其内部结构是嵌套的字典和列表。
    :param columns_num: 一个整数值，表示在开始提取信息之前要跳过的列数。

    函数内部逻辑如下：
    """
    # 用于存储单元格坐标信息的列表，每个元素是一个包含坐标点元组的列表
    rectangles = []
    # 用于存储tableCellId信息的列表
    tableCellIds = []
    # 用于存储行数信息的列表
    num = []

    count = 0
    st_ysc = -1
    st_num = -1
    temp_points = []

    # 遍历prism_tables_info中的每个表格信息
    for table_info in prism_tables_info:
        # 遍历每个表格中的单元格信息
        for cell_info in table_info.get('cellInfos', []):
            # 如果计数大于等于要跳过的列数加1，则开始提取信息
            if count >= columns_num + 1:
                # 获取单元格的位置信息（字典），如果不存在则返回空字典
                pos = cell_info.get('pos', {})

                # 遍历位置信息中的每个点
                for point in pos:
                    # 获取x坐标
                    x = point['x']
                    # 获取y坐标
                    y = point['y']

                    # 将坐标点添加到临时列表中
                    temp_points.append((x, y))
                # 将临时坐标点列表添加到rectangles列表中
                rectangles.append(temp_points)

                # 获取单元格的tableCellId，如果不存在则默认为0，然后转换为整数类型并添加到tableCellIds列表中
                tableCellId = int(cell_info.get('tableCellId', 0))
                tableCellIds.append(tableCellId)

                # 获取单元格的ysc值，如果不存在则默认为0
                ysc = cell_info.get('ysc', 0)
                # 获取单元格的word值，如果不存在则默认为0
                word = cell_info.get('word', 0)

                # 如果当前ysc值大于之前的最大ysc值
                if ysc > st_ysc:
                    st_ysc = ysc
                    st_num = word
                # 将当前最大行数添加到num列表中
                num.append(st_num)

                # 清空临时坐标点列表，为下一个单元格做准备
                temp_points = []
            count += 1

    # 使用提取到的信息构建一个字典
    data = {
        'num': num,
        'tableCellId': tableCellIds,
        'axes': rectangles
    }

    # 使用字典创建一个DataFrame
    df = pd.DataFrame(data)

    # 在DataFrame的第一列插入原始图像路径列
    df.insert(loc=0, column='orgImage_path', value=org_image_path)

    # 重新排列DataFrame的列顺序
    columns = ['orgImage_path', 'num', 'tableCellId', 'axes']
    df = df.reindex(columns=columns)

    # 将DataFrame保存到Excel文件中，不包含索引列
    df.to_excel(excel_file_name, index=False)

    # 返回1，表示函数执行成功（可能是一种简单的成功标识方式）
    return 1

def tableTitleRowNum(json_data):
    """
    该函数的主要功能是从给定的JSON数据中获取与表格标题行相关的信息，包括标题行的数量、
    由tableCellID和word组成的元组列表（这里代码中实际上只构建了包含word的列表，与函数描述部分不完全一致）以及标题行tableCellID的数量（这里实际是标题行所有字段的粗略数量）。

    :param json_data: 输入的JSON数据，数据结构是嵌套的字典形式，其中包含了表格相关的信息。例如：{"body": {"Data": "具体的表格数据信息的字符串表示"}}。
    :return: 返回一个包含三个元素的元组，分别为标题行数量、除去标题行的数据条目数量、标题所有字段的粗略数量。

    函数内部逻辑如下：
    """
    # 首先对json_data["body"]["Data"]进行eval操作，将其转换为合适的Python对象（这里假设是字典），并获取其中的"prism_tablesInfo"部分
    json_data1 = eval(json_data["body"]["Data"])
    prism_tables_info = json_data1["prism_tablesInfo"]

    titleRowNum = 0
    titleRow = []

    # 外层循环遍历prism_tables_info中的每个表格信息
    for table_info in prism_tables_info:
        # 获取表格的x方向单元格大小和y方向单元格大小
        xCellSize = table_info['xCellSize']
        yCellSize = table_info['yCellSize']
        tableTitleLen = xCellSize

        # 内层循环遍历每个表格中的单元格信息
        for cell_info in table_info.get('cellInfos', []):
            # 获取单元格的yec、ysc和xsc值，如果不存在则默认为0
            yec0 = cell_info.get('yec', 0)
            ysc0 = cell_info.get('ysc', 0)
            xsc0 = cell_info.get('xsc', 0)
            # 如果xsc0小于表格标题长度
            if xsc0 < tableTitleLen:
                # 计算行的大小（这里的计算逻辑基于yec0和ysc0，具体含义取决于数据结构的定义）
                row0 = yec0 - ysc0 + 1
                # 如果计算得到的行大小大于当前的标题行数量
                if row0 > titleRowNum:
                    titleRowNum = row0

    # 再次遍历prism_tables_info中的每个表格信息
    for table_info in prism_tables_info:
        for cell_info in table_info.get('cellInfos', []):
            # 获取单元格的ysc和word值，如果不存在则默认为0
            ysc0 = cell_info.get('ysc', 0)
            word0 = cell_info.get('word', 0)
            # 如果ysc0小于标题行数量
            if ysc0 < titleRowNum:
                # 将word0添加到titleRow列表中
                titleRow.append(word0)

    titleRowLen = len(titleRow)
    dataRowNum = yCellSize - titleRowNum

    # 返回标题行数量、除去标题行的数据条目数量、标题所有字段的粗略数量
    return titleRowNum, dataRowNum, titleRowLen

def tableTitle(json_data):
    """
    该函数主要从给定的JSON数据中提取表格标题相关的有效字段名称等信息。

    :param json_data: 输入的JSON数据，其结构为嵌套的字典形式，例如{"body": {"Data": "具体的数据字符串"}}，其中包含了表格相关的信息。
    :return: 返回一个包含三个元素的元组，分别是列表形式的标题有效字段名称、标题所有字段的粗略数量、除去标题行的数据条目数量。

    函数内部逻辑如下：
    """
    # 首先对json_data["body"]["Data"]进行eval操作，将其转换为合适的Python对象（这里假设是字典），并获取其中的"prism_tablesInfo"部分
    json_data1 = eval(json_data["body"]["Data"])
    prism_tables_info = json_data1["prism_tablesInfo"]

    temp = []
    temp0 = []
    row = []

    # 调用tableTitleRowNum函数获取标题行数量、除去标题行的数据条目数量、标题所有字段的粗略数量，这里只使用后两个返回值
    _, dataRowNum, titleRowLen = tableTitleRowNum(json_data)

    # 遍历prism_tables_info中的每个表格信息
    for table_info in prism_tables_info:
        # 遍历每个表格中前titleRowLen个单元格信息（这里假设是标题行的单元格）
        for cell_info in table_info.get('cellInfos', [])[:titleRowLen]:
            # 获取单元格的yec、ysc、xec、xsc和word值，如果不存在则默认为0
            yec0 = cell_info.get('yec', 0)
            ysc0 = cell_info.get('ysc', 0)
            xec0 = cell_info.get('xec', 0)
            xsc0 = cell_info.get('xsc', 0)
            word0 = cell_info.get('word', 0)
            # 计算x方向和y方向的单元格跨度
            xse = xec0 - xsc0
            yse = yec0 - ysc0

            # 如果ysc0为0，将word0添加到row列表中，这里可能是将同一行（ysc0为0表示行首）的word值添加到row中
            if ysc0 == 0:
                row.append(word0)
            # 如果x方向跨度大于0且ysc0为0且y方向跨度为0，将相关信息（xsc0、xse、word0）添加到temp列表中
            if xse > 0 and ysc0 == 0 and yse == 0:
                temp.append((xsc0, xse, word0))
            # 如果x方向跨度为0且ysc0为1且y方向跨度为0，将word0添加到temp0列表中
            if xse == 0 and ysc0 == 1 and yse == 0:
                temp0.append(word0)

    # 当temp列表不为空时进行循环操作
    while temp:
        # 弹出temp列表的第一个元素，并获取这个元素
        pop_element = temp.pop(0)
        # 根据pop_element的第二个元素（xse的值）确定从temp0中提取的元素数量，并获取这些元素
        extracted_values = temp0[:pop_element[1] + 1]
        # 在row列表中查找pop_element的第三个元素（word0）的索引
        index = row.index(pop_element[2])
        # 根据索引从row列表中删除对应的元素
        row.pop(index)
        # 获取row列表中索引之前的部分
        left_part = row[:index]
        # 获取row列表中索引之后的部分
        right_part = row[index:]
        # 将left_part、extracted_values和right_part合并成一个新的列表
        my_list = left_part + extracted_values + right_part
        # 遍历extracted_values中的每个元素，在temp0列表中查找并删除该元素
        for value in extracted_values:
            temp0.pop(temp0.index(value))

    # 返回处理后的列表形式的标题有效字段名称、标题所有字段的粗略数量、除去标题行的数据条目数量
    return my_list, titleRowLen, dataRowNum

def json_to_df(json_data, db_query, filepath, key_table_name="xschemeKeyTable", dfType="list-dict"):
    """
    该函数的主要功能是从给定的JSON数据、数据库查询对象和文件路径中提取相关信息，构建一个DataFrame，并根据指定的类型返回相应的数据结构，同时还返回一些与表格结构相关的统计信息。

    :param json_data: 输入的JSON数据，其结构为嵌套的字典形式，例如{"body": {"Data": "具体的数据字符串"}}，包含了表格相关的信息。
    :param db_query: 数据库查询对象，用于在数据库中根据特定条件查询数据，具体的查询逻辑取决于该对象的实现。
    :param filepath: 文件路径，字符串类型，用于从路径中提取特定的代码（如archiveCode）。
    :param key_table_name: 数据库中用于查询的表名，默认值为"xschemeKeyTable"，字符串类型。
    :param dfType: 表示返回数据的类型，可选值为"list-dict"或"df"，默认值为"list-dict"，字符串类型。

    :return: 返回一个包含五个元素的元组，分别是根据dfType处理后的表格数据（df形式或list-dict形式）、
             一个包含所有单元格字段值都相同的字段的字典（如文件路径、年份、省份等）、
             标题行的粗略长度、数据条目的长度、表格的单元格数量。

    函数内部逻辑如下：
    """
    # 从JSON数据中获取 "Data" 字段的值，并进一步获取 "prism_tablesInfo" 和 "tableHeadTail" 部分
    json_data1 = eval(json_data["body"]["Data"])
    prism_tables_info = json_data1["prism_tablesInfo"]
    tableHeadTail = json_data1["tableHeadTail"]
    tableHead0 = tableHeadTail[0]['head'][0]
    tableHead1 = tableHeadTail[0]['head'][1]

    st_ysc = -1
    st_num = -1

    countToll = 0
    temp_points = []

    tableCellId = []
    yec = []
    ysc = []
    xec = []
    xsc = []
    num = []
    archiveCode = []
    word = []
    axes = []
    keyCode = []

    # 定义正则表达式来匹配四位数字作为年份
    year_pattern = r"(\d{4})"
    # 使用正则表达式在tableHead0中进行匹配，获取匹配到的年份
    year = re.search(year_pattern, tableHead0)[0]

    # 定义正则表达式来匹配省份名称
    province_pattern = r'(?:(北京|天津|上海|重庆)|(河北|山西|内蒙古|辽宁|吉林|黑龙江|江苏|浙江|安徽|福建|江西|山东|河南|湖北|湖南|广东|海南|四川|贵州|云南|陕西|甘肃|青海|台湾)|(广西|西藏|宁夏|新疆)|(香港|澳门))'
    # 使用正则表达式在tableHead0中进行匹配，获取匹配到的省份名称
    province = re.search(province_pattern, tableHead0)[0]

    # 定义正则表达式来匹配招生批次相关的字符串
    admission_batch_pattern = r'([本科|专科|研究生|普通类本科批|特殊类型|提前批|第一批|文理本科|文理科本科|国家|高校]+)(普通批|提前批|第一批|一批|二批|高职\(专科\)批|A阶段|招生志愿|艺术B段|第一批B段|提前B|专项)'
    # 使用正则表达式在tableHead1中进行匹配，获取匹配到的招生批次字符串
    admissionBatch = re.search(admission_batch_pattern, tableHead1)[0]

    # 定义正则表达式来匹配特定格式（1950 - JX11.44 - 10形式）的字符，用于从文件路径中获取archiveCode
    archive_code_pattern = r'\d{4}-[A-Za-z0-9]{2}\d{2}\.[A-Za-z0-9]{2}-\d+'
    # 使用正则表达式在filepath中进行匹配，获取匹配到的archiveCode
    archiveCode = re.search(archive_code_pattern, filepath)[0]

    # 调用tableTitle函数获取标题相关信息，包括标题有效字段名称、标题所有字段粗略数量、除去标题行的数据条目数量
    titleList, titleRowLen, dataRowNum = tableTitle(json_data)

    # 遍历prism_tables_info中的每个表格信息
    for table_info in prism_tables_info:
        xCellSize = table_info['xCellSize']
        yCellSize = table_info['yCellSize']
        for cell_info in table_info.get('cellInfos', []):
            # 如果计数大于等于标题行的粗略长度，则开始提取单元格相关信息
            if countToll >= titleRowLen:
                # 获取单元格的位置信息（字典形式），如果不存在则默认为空列表
                pos = cell_info.get('pos', [])
                for point in pos:
                    x = point['x']
                    y = point['y']
                    temp_points.append((x, y))
                # 将临时坐标点列表添加到axes列表中，表示单元格的坐标信息
                axes.append(temp_points)
                temp_points = []

                # 获取单元格的tableCellId，如果不存在则默认为0，转换为整数类型后添加到tableCellId列表中
                tableCellId.append(int(cell_info.get('tableCellId', 0)))
                word0 = cell_info.get('word', 0)
                word.append(word0)
                yec0 = cell_info.get('yec', 0)
                yec.append(yec0)
                ysc0 = cell_info.get('ysc', 0)
                ysc.append(ysc0)
                xec0 = cell_info.get('xec', 0)
                xec.append(xec0)
                xsc0 = cell_info.get('xsc', 0)
                xsc.append(xsc0)

                # 如果当前ysc0大于之前的最大ysc值
                if ysc0 > st_ysc:
                    st_ysc = ysc0
                    st_num = word0
                num.append(st_num)

                # 根据标题列表和当前单元格的xsc0获取查询条件中的键名
                keyname_to_query = titleList[xsc0]
                # 构建查询条件，用于在数据库中查询keyCode
                condition = f"FIND_IN_SET('{keyname_to_query}', REPLACE(orgName, '|', ',')) > 0"

                # 使用数据库查询对象根据条件查询keyCode
                query_result = db_query.query_data(key_table_name, select_fields='keyCode', condition=condition)
                default_value = "NA"

                # 如果查询结果存在，则将查询到的第一个keyCode添加到keyCode列表中，否则添加默认值"NA"
                if query_result:
                    keyCode.append(query_result[0][0])
                else:
                    keyCode.append(default_value)

            countToll += 1

    # 使用提取到的各种信息构建一个字典
    data = {
        'yec': yec,
        'ysc': ysc,
        'xec': xec,
        'xsc': xsc,
        'word': word,
        'num': num,
        'tableCellId': tableCellId,
        'axes': axes,

        'archiveCode': archiveCode,
        'keyCode': keyCode,
        'year': year,
        'province': province,
        'admissionBatch': admissionBatch
    }

    # 使用字典创建一个DataFrame
    df = pd.DataFrame(data)

    # 在DataFrame的第一列插入文件路径列
    df.insert(loc=0, column='filepath', value=filepath)
    # 在DataFrame中添加xCellSize和yCellSize列
    df['xCellSize'] = xCellSize
    df['yCellSize'] = yCellSize

    # 构建一个包含所有单元格字段值都相同的字段的字典，如文件路径、archiveCode、年份、省份、招生批次等
    commonDict = {"filePath": filepath, "archiveCode": archiveCode, "year": year, "province": province,
                  "admissionBatch": admissionBatch}

    # 根据dfType的值处理要返回的数据
    if dfType == "list-dict":
        data_list = df.to_dict(orient='records')
        result = data_list
    elif dfType == "df":
        result = df

    # 计算表格的单元格数量
    cellNum = xCellSize * dataRowNum

    # 返回处理后的表格数据、公共字段字典、标题行粗略长度、数据条目长度、单元格数量
    return result, commonDict, titleRowLen, dataRowNum, cellNum

def df_db(table_name, cell_name, db_query, json_data, filepath, logName='image2tableProcessLog'):
    """
    该函数的主要功能是将从JSON数据转换得到的DataFrame数据存入数据库，并根据不同的情况进行日志记录和数据处理。

    :param table_name: 数据库中的表名，字符串类型，用于确定将数据插入到哪个表中。
    :param cell_name: 数据库中的单元格表名称（与存储单元格数据相关的标识），字符串类型。
    :param db_query: 数据库查询对象，用于执行数据库的插入、查询、更新、删除等操作，具体的操作逻辑取决于该对象的实现。
    :param json_data: 输入的JSON数据，其结构为嵌套的字典形式，包含了表格相关的信息，将被转换为DataFrame并插入数据库。
    :param filepath: 文件路径，字符串类型，用于在数据库查询中构建查询条件等操作。
    :param logName: 日志表名称，默认值为'image2tableProcessLog'，字符串类型，用于在数据库中记录操作相关的日志信息。

    """
    # 调用json_to_df函数将JSON数据转换为DataFrame，并获取相关的统计信息
    result, commonDict, titleRowLen, dataRowNum, cellNum = json_to_df(json_data, db_query, filepath=filepath, dfType='df')

    # 构建查询条件，用于查询数据库中特定文件路径对应的记录数量
    select_fields = f"COUNT(*)"
    condition = f"filePath='{filepath}'"
    # 使用数据库查询对象查询特定表中满足条件的记录数量
    checkTable = db_query.query_data(table_name, select_fields, condition=condition)[0][0]

    # 获取从JSON转换得到的DataFrame的长度（记录数量）
    resultLen = len(result)

    # 尝试将DataFrame中的数据插入到数据库中（以字典形式的记录），如果插入成功则根据不同情况进行处理
    if (db_query.insert_row(cell_name, result.to_dict(orient='records'))):
        # 如果插入的记录数量等于计算得到的单元格数量
        if resultLen == cellNum:
            # 打印日志，表示OCR识别完整，JSON已录入cell数据库
            print(f"'{filepath}' OCR识别完整，json已录入cell数据库;")
            value = {'cell': 1}
            # 使用数据库查询对象更新日志表中的记录
            db_query.update_row(logName, value, condition)
        # 如果插入的记录数量大于0且小于计算得到的单元格数量
        elif 0 < resultLen < cellNum:
            value = " \\\\n " + str(datetime.datetime.now()) + " OCR识别单元格Cell不完整，json已录入Cell;"
            # 打印日志
            print(value)
            # 使用数据库查询对象向日志表的errorInfo字段追加值
            db_query.append_value(logName, "errorInfo", value, condition)
            value = {'cell': 2}
            # 使用数据库查询对象更新日志表中的记录
            db_query.update_row(logName, value, condition)
        else:
            value = " \\\\n " + str(datetime.datetime.now()) + " OCR识别单元格Cell失败或超出可能数量，将删除已导入Cell单元格数据;"
            # 使用数据库查询对象删除特定条件下的单元格数据
            db_query.delete_row(cell_name, condition)
            # 打印日志
            print(value)
            # 使用数据库查询对象向日志表的errorInfo字段追加值
            db_query.append_value(logName, "errorInfo", value, condition)

    # 根据查询到的特定表中的记录数量与数据条目的数量关系进行不同的处理
    if checkTable == dataRowNum:
        value = " \\\\n " + str(datetime.datetime.now()) + " json已录入table数据库;"
        # 使用数据库查询对象向日志表的errorInfo字段追加值
        db_query.append_value(logName, "errorInfo", value, condition)
        value = {'table': 1}
        # 使用数据库查询对象更新日志表中的记录
        db_query.update_row(logName, value, condition)
    elif 0 < checkTable < dataRowNum:
        value = " \\\\n " + str(datetime.datetime.now()) + " json录入table数据时发生问题，将删除录入不完整的数据;"
        # 使用数据库查询对象向日志表的errorInfo字段追加值
        db_query.append_value(logName, "errorInfo", value, condition)
        # 使用数据库查询对象删除特定条件下的表数据
        db_query.delete_row(table_name, condition)
        value = {'tab': 0}
        # 使用数据库查询对象更新日志表中的记录
        db_query.update_row(logName, value, condition)
    else:
        # 将DataFrame中的'word'列转换为列表
        word_list = result['word'].tolist()
        # 将DataFrame中的'ysc'列转换为列表
        ysc_list = result['ysc'].tolist()
        # 每titleRowLen - 1行为一组，创建新的DataFrame
        x = titleRowLen - 1
        new_data = []
        current_num = None
        row_data = []

        # 遍历'word'和'ysc'列表，根据ysc的值对数据进行分组处理
        for word, ysc in zip(word_list, ysc_list):
            if ysc!= current_num:
                if current_num is not None:
                    # 如果当前行数据不足x个，则用'NA'填充
                    row_data += ['NA'] * (x - len(row_data))
                    new_data.append(row_data)

                current_num = ysc
                row_data = []

            row_data.append(word)

        # 处理最后一行数据，用'NA'填充不足的部分
        row_data += ['NA'] * (x - len(row_data))
        new_data.append(row_data)
        # 获取'keyCode'列的唯一值列表
        keycode_list = result['keyCode'].unique()
        # 使用处理后的数据创建新的DataFrame，并转换为字典形式（以记录为导向）
        df_dict = pd.DataFrame(new_data, columns=keycode_list).to_dict(orient='records')

        # 遍历新的字典列表，更新每个字典元素并尝试插入到数据库表中
        for d in df_dict:
            d.update(commonDict)
            if (db_query.insert_row(table_name, df_dict)):
                value = {'tab': 1}
                # 使用数据库查询对象更新日志表中的记录
                db_query.update_row(logName, value, condition)

def format_image_from_excel(orgImage_path, excel_file, row_number):
    """
    该函数的主要功能是从包含cell信息的Excel表中获取坐标信息，然后在原始图像上对除了给定行之外的行进行填充操作，最后输出与原始图像质量相同的图像。
    坐标信息格式：[(x1, y1), (x2, y2),(x3, y3),(x4, y4)]
    :param orgImage_path: 原始图像的路径，字符串类型，用于打开原始图像进行填充操作。
    :param excel_file: Excel文件的路径，字符串类型，该文件包含了用于填充操作的坐标等相关信息。
    :param row_number: 一个整数值，表示需要排除填充的行号，即对除了这一行之外的其他行进行填充操作。

    函数内部逻辑如下：
    """
    # 使用pandas的read_excel函数从Excel文件中读取数据，并将其存储为DataFrame
    df = pd.read_excel(excel_file)

    # 使用isin方法检查DataFrame中的'num'列是否包含给定的行号（row_number），如果包含则执行后续操作
    if df['num'].isin([row_number]).any():
        # 使用PIL库的Image.open函数打开原始图像，得到一个Image对象
        image = Image.open(orgImage_path)

        # 从图像的信息中获取分辨率（dpi），如果图像信息中不存在dpi，则返回None
        dpi = image.info.get('dpi')
        # 创建一个可以在图像上进行绘制操作的对象
        draw = ImageDraw.Draw(image)
        rectangle_points_list = []

        # 遍历DataFrame的每一行，行索引为index，行数据为row
        for index, row in df.iterrows():
            # 如果当前行的'num'值不等于给定的行号（row_number）
            if row['num']!= row_number:
                # 将当前行的'axes'列的值添加到rectangle_points_list中，这里的'axes'列可能包含了用于绘制填充区域的坐标信息
                rectangle_points_list.append(row['axes'])

        # 遍历rectangle_points_list中的每个坐标列表
        for rectangle_points in rectangle_points_list:
            # 使用eval函数将坐标字符串转换为实际的坐标值（这里假设坐标是以字符串形式存储在Excel中的），然后使用draw.polygon方法绘制多边形并进行填充
            # 多边形的轮廓颜色为黑色，填充颜色为白色
            draw.polygon(eval(rectangle_points), outline='black', fill='white')

        # 定义输出图像的保存路径
        output_path = r"./0001test.jpg"
        # 使用image.save方法将填充后的图像保存到指定路径，保存格式为JPEG，并设置分辨率为原始图像的dpi
        image.save(output_path, 'JPEG', dpi=dpi)
        print("图片已保存。")
    else:
        print("数据不存在.")

def format_image_from_db(table, orgImage_path, db_query, row_number, output_path=''):
    """
    该函数的主要功能是从数据库中的表格获取坐标信息，然后在原始图像上对除给定行之外的行进行填充操作，最后输出与原始图像质量相同的图像（可指定输出路径）。
    坐标信息格式：[(x1, y1), (x2, y2),(x3, y3),(x4, y4)]
    :param table: 数据库中的表名，字符串类型，用于指定从哪个表中查询数据。
    :param orgImage_path: 原始图像的路径，字符串类型，用于打开原始图像进行填充操作。
    :param db_query: 数据库查询对象，用于执行数据库查询操作，具体的查询逻辑取决于该对象的实现。
    :param row_number: 一个整数值，表示需要排除填充的行号，即对除了这一行之外的其他行进行填充操作。
    :param output_path: 可选参数，输出图像的路径，字符串类型。如果未提供，函数将返回处理后的图像对象但不保存图像；如果提供了路径，则将填充后的图像保存到该路径下。

    函数内部逻辑如下：
    """
    # 构建查询条件，用于从数据库中查询符合要求的数据。要求是文件路径等于原始图像路径并且行号不等于给定的行号
    condition = f"{'filePath'} = '{orgImage_path}' AND {'num'}!= '{row_number}'"

    # 使用数据库查询对象查询指定表中的'axes'字段，根据构建的条件进行查询
    result = db_query.query_data(table, 'axes', condition)
    if result:
        # 如果查询结果存在，将查询结果中的每个子列表的第一个元素（假设每个结果行是一个包含单个元素的列表，该元素为坐标信息）提取出来，组成新的列表
        result = [row[0] for row in result]
    else:
        # 如果查询结果为空，直接返回空列表
        return []

    # 使用PIL库的Image.open函数打开原始图像，得到一个Image对象
    image = Image.open(orgImage_path)
    # 从图像的信息中获取分辨率（dpi），如果图像信息中不存在dpi，则返回None
    dpi = image.info.get('dpi')
    # 创建一个可以在图像上进行绘制操作的对象
    draw = ImageDraw.Draw(image)

    # 遍历查询得到的坐标列表
    for rectangle_points in result:
        # 使用eval函数将坐标字符串转换为实际的坐标值（这里假设坐标是以字符串形式存储在数据库中的），然后使用draw.polygon方法绘制多边形并进行填充
        # 多边形的轮廓颜色为黑色，填充颜色为白色
        draw.polygon(eval(rectangle_points), outline='black', fill='white')

    # 如果提供了输出路径
    if output_path:
        # 使用image.save方法将填充后的图像保存到指定路径，保存格式为JPEG，并设置分辨率为原始图像的dpi
        image.save(output_path, 'JPEG', dpi=dpi)
    # 如果未提供输出路径，返回处理后的图像对象
    return image


def check_image_resolution(img, max_width, max_height):
    """
    该函数用于检查图像在横向和纵向两个方向的分辨率是否超过指定的阈值。

    :param img: 要检查的图像对象，是通过PIL库打开后的图像对象。
    :param max_width: 最大宽度阈值，整数类型，表示图像宽度的上限。
    :param max_height: 最大高度阈值，整数类型，表示图像高度的上限。
    :return: 如果图像的宽度和高度都小于等于对应的阈值，则返回True，表示未超过阈值；否则返回False，表示超过阈值；如果在获取图像尺寸过程中发生异常，则返回None。
    """
    print("check_image_resolution 检查分辨率")
    try:
        # 使用图像对象的size属性获取图像的宽度和高度
        width, height = img.size

        # 检查图像的宽度是否小于等于最大宽度阈值，并且高度是否小于等于最大高度阈值
        if width <= max_width and height <= max_height:
            return True
        else:
            return False
    except Exception as e:
        # 如果在获取图像尺寸过程中发生异常，打印异常信息（这里假设print函数是用于打印日志的自定义函数）
        print("check_image_resolution Error:", e)
        return None


def resize_image_resolution(img, max_size):
    """
    该函数的主要功能是按比例缩小图像，确保图像长宽的最大值不超过指定的数值。

    :param img: 要处理的图像对象，是通过PIL库打开后的图像对象。
    :param max_size: 缩小后的长宽最大值，整数类型，表示图像缩小后的最大边长（宽度或高度）。
    :return: 如果成功缩小图像，则返回缩小后的图像对象；如果在处理过程中发生异常，则返回None。
    """
    try:
        # 获取图像的宽度和高度
        width, height = img.size
        # 计算图像的最大边长（宽度和高度中的较大值）
        max_dimension = max(width, height)
        # 计算缩小比例，即当前最大边长与指定最大边长的比值
        scale_factor = max_dimension / max_size

        # 根据缩小比例计算缩小后的宽度和高度
        new_width = int(width / scale_factor)
        new_height = int(height / scale_factor)

        # 使用图像对象的thumbnail方法进行等比缩放，将图像缩小到计算得到的新宽度和新高度
        img.thumbnail((new_width, new_height))

        # 打印日志，表示图像已成功按比例缩小（这里假设print函数是用于打印日志的自定义函数）
        print("图像已成功按比例缩小，长宽的最大值不超过指定数值")
        return img
    except Exception as e:
        # 如果在处理过程中发生异常，打印异常信息（这里假设print函数是用于打印日志的自定义函数）
        print("Error:", e)
        return None

def delete_file(file_path):
    try:
        os.remove(file_path)
        print(f"{file_path} deleted successfully")
    except OSError as e:
        print(f"Error deleting {file_path}: {e}")
        
def delete_files_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
            elif os.path.isdir(file_path):
                delete_files_in_folder(file_path)  # 递归调用，删除子文件夹中的文件
        except Exception as e:
            print(f"Failed to delete {file_path}: {e}")
def delete_items(item_list):
    """
    删除给定列表中的文件或目录。

    参数：
    - item_list：一个包含要删除的文件或目录路径的列表。

    功能描述：
    1. 遍历列表中的每个路径。
    2. 检查路径是否存在，如果不存在则打印错误信息并跳过当前项。
    3. 如果是文件，则直接删除。
    4. 如果是目录，则递归地删除目录及其下所有子目录和文件。
    """
    for item_path in item_list:
        if not os.path.exists(item_path):
            print(f"路径 '{item_path}' 不存在。")
            continue

        if os.path.isfile(item_path):
            os.remove(item_path)
            print(f"文件 '{item_path}' 已被删除。")
        elif os.path.isdir(item_path):
            shutil.rmtree(item_path)
            print(f"目录 '{item_path}' 及其内容已被删除。")
        else:
            print(f"路径 '{item_path}' 不是文件或目录。")
            
def move_files_or_directories(source_target_mapping):  
    """  
    将源目录或文件移动到目标路径。  

    参数:  
    - source_target_mapping: 一个字典，键是源路径（source_path），值是目标路径（target_path）。  

    功能描述:  
    1. 遍历字典中的每个键值对，键是源路径，值是目标路径。  
    2. 检查源路径是否存在，如果不存在则打印错误信息并跳过当前项。  
    3. 检查源路径是文件还是目录。  
    4. 如果是文件，则直接移动到目标路径。  
    5. 如果是目录，则递归地移动目录和其下所有子目录和文件到目标路径。  
    6. 进行必要的容错处理，如目标路径已存在时的处理。  
    """  
    for source_path, target_path in source_target_mapping.items():  
        # 检查源路径是否存在  
        if not os.path.exists(source_path):  
            print(f"源路径 '{source_path}' 不存在。")  
            continue  

        # 检查源路径是文件还是目录  
        if os.path.isfile(source_path):  
            # 源路径是文件，直接移动文件  
            shutil.move(source_path, target_path)  
            print(f"文件 '{source_path}' 已移动到 '{target_path}'。")  
        elif os.path.isdir(source_path):  
            # 源路径是目录，移动目录及其所有子目录和文件  
            # 确保目标路径的父目录存在  
            target_dir = os.path.dirname(target_path)  
            if target_dir and not os.path.exists(target_dir):  
                os.makedirs(target_dir)  # 创建目标路径中的目录（如果不存在）  
            
            if os.path.exists(target_path):  
                print(f"目标路径 '{target_path}' 已存在，将直接覆盖该目录。")  
            else:  
                os.makedirs(target_path, exist_ok=True)  # 创建目标目录（如果不存在）  
            
            shutil.move(source_path, target_path)  
            print(f"目录 '{source_path}' 及其内容已移动到 '{target_path}'。")  
        else:  
            print(f"路径 '{source_path}' 不是文件或目录。")  
    return True
    
def count_files_in_folder(folder_path):
    # 初始化文件数量
    file_count = 0

    # 遍历文件夹内所有文件和子文件夹
    for root, dirs, files in os.walk(folder_path):
        file_count += len(files)

    return file_count
    
# 将变量以字典的形式{"变量":"值"}存入文本文件，如果值存在就覆盖，不存在就追加,file_path是存储变量的文件名
def write_variables_to_file(variables_dict,file_path):
    existing_variables = {}
    # 读取文件中已有的变量
    if os.path.exists(file_path):
        with open(file_path, 'r') as file:
            for line in file:
                key, value = line.strip().split('=')
                existing_variables[key] = value
    # 更新或追加变量
    with open(file_path, 'w') as file:
        for key, value in variables_dict.items():
            existing_variables[key] = value
        # 写入更新后的所有变量
        for key, value in existing_variables.items():
            file.write(f"{key}={value}\n")

# 从文件中读取变量，如果文件或变量不存在就返回default设定的值,file_path是存储变量的文件名
def read_variable_from_file(default,variable_name,file_path):
    if not os.path.exists(file_path):
        with open(file_path, 'w') as file:
            pass  # 'pass' 表示我们不做任何操作，这里只是创建一个空文件
        return default  # 文件不存在，返回default
    with open(file_path, 'r') as file:
        for line in file:
            key, value = line.strip().split('=')
            if key == variable_name:
                return value
    return default  # 变量不存在，返回默认值

# 从前往后截取字符串，前面用...代替
def truncate_string(input_string, max_length):
    if len(input_string) > max_length:
        return "..." + input_string[-max_length+3:]
    else:
        return input_string

def save_settings(file_path, settings_list,messagebox_flag=True,parent=None,):
    """
    保存设置到指定的文件路径。
    参数:
    - file_path: 要保存的文件路径。
    - settings_list: 要保存的参数列表，格式为 [key=value,...]。
    - parent: 消息框的父窗口，默认为 None。
    - messagebox_flag : True显示成功对话框，False不显示。
    """
    try:
        # 确保目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        existing_settings = {}
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                lines = f.readlines()
                for line in lines[1:]:
                    parts = line.strip().split('=')
                    if len(parts) == 2:
                        existing_settings[parts[0]] = parts[1]

        new_settings = {}
        for setting in settings_list:
            parts = setting.split('=')
            if len(parts) == 2:
                new_settings[parts[0]] = parts[1]

        combined_settings = {**existing_settings, **new_settings}

        with open(file_path, 'w') as f:
            f.write(f'[Settings]\n')
            for key, value in combined_settings.items():
                f.write(f"{key}={value}\n")

        if messagebox_flag==True:
            # 如果保存成功，弹出消息框告知用户
            print(f"信息已成功保存到 {file_path} 文件。")
            # QMessageBox.information(parent, "成功", f"信息已成功保存到 {file_path} 文件。")

    except Exception as e:
        if messagebox_flag==True:
            # 如果发生错误，弹出消息框告知用户错误信息
            print(f"无法保存信息到 {file_path} 文件。错误: {e}")
            # QMessageBox.critical(parent, "错误", f"无法保存信息到 {file_path} 文件。错误: {e}")

def read_settings(file_path, variable_names):
    """
    从指定的 INI 文件中批量读取设置。
    参数:
    - file_path: 要读取的文件路径。
    - variable_names: 要读取的变量名称列表。
    """
    settings = {}
    # 创建 ConfigParser 对象
    config = configparser.ConfigParser()

    # 检查文件是否存在
    if not os.path.exists(file_path):
        print(f"Error: The file {file_path} does not exist.")
        return settings

    # 读取 INI 文件
    config.read(file_path)

    # 默认 section 名称
    section = 'Settings'

    # 检查 section 是否存在
    if not config.has_section(section):
        print(f"Error: The section {section} does not exist in the file.")
        return settings

    for variable_name in variable_names:
        # 检查变量是否存在
        if not config.has_option(section, variable_name):
            print(f"Error: The variable {variable_name} does not exist in the section {section}.")
            settings[variable_name] = None
        else:
            # 获取变量值
            value = config.get(section, variable_name)
            settings[variable_name] = value

    return settings

def read_setting(file_path, variable_name):  
    """  
    从指定的 INI 文件中读取设置。  
    参数:  
    - file_path: 要读取的文件路径。  
    - variable_name: 要读取的变量名称。  
    """  
    # 创建 ConfigParser 对象  
    config = configparser.ConfigParser()  
    
    # 检查文件是否存在  
    if not os.path.exists(file_path):  
        print(f"Error: The file {file_path} does not exist.")  
        return None  
    
    # 读取 INI 文件  
    config.read(file_path)  
    
    # 默认 section 名称  
    section = 'Settings'  
    
    # 检查 section 是否存在  
    if not config.has_section(section):  
        print(f"Error: The section {section} does not exist in the file.")  
        return None  
    
    # 检查变量是否存在  
    if not config.has_option(section, variable_name):  
        print(f"Error: The variable {variable_name} does not exist in the section {section}.")  
        return None  

    # 获取变量值  
    value = config.get(section, variable_name)  
    return value  

def insert_empty_rows(df, records_count, file_start, file_column_name, record_column_name):  
    """  
    在给定的 DataFrame 中插入空行，并为指定列赋值。  

    参数:  
    - df: 输入的 Pandas DataFrame，包含需要处理的数据。  
    - records_count: 每组记录的数量，用于确定插入空行的位置。  
    - file_start: 用于生成文件编号的起始值。  
    - file_column_name: 新增列的名称，用于存储文件编号。  
    - record_column_name: 新增列的名称，用于存储记录编号。  

    返回:  
    - 返回一个新的 DataFrame，其中包含插入的空行和新的编号列。  
    """  
    new_rows = []  
    # 如果 records_count 是字符串，尝试转换为整数
    if isinstance(records_count, str):
        records_count = int(records_count)
    # 插入一行空行作为起始行  
    new_rows.append([None] * len(df.columns))  
    
    row_index = 0  
    
    # 遍历原 DataFrame 的每一行  
    while row_index < len(df):  
        # 将当前行添加到新行列表中  
        new_rows.append(df.iloc[row_index].tolist())  
        
        # 每 records_count 行插入一个空行  
        if (row_index + 1) % records_count == 0:  
            new_rows.append([None] * len(df.columns))  # 插入空行  
            
        row_index += 1  # 移动到下一行  

    # 创建新的 DataFrame  
    new_df = pd.DataFrame(new_rows, columns=df.columns)  
    
    # 如果最后一行全为 None，则删除最后一行  
    if all(new_df.iloc[-1].isnull()):  
        new_df = new_df.iloc[:-1]  
    
    file_column = []  # 用于存储文件编号  
    row_num = new_df.shape[0]  # 获取新 DataFrame 的行数  
    current_start = file_start  # 初始化文件编号的起始值  
    
    # 为每组记录生成文件编号  
    for _ in range(row_num // records_count + 2):  
        for _ in range(records_count + 1):  
            if len(file_column) < row_num:  
                file_column.append(current_start)  # 添加当前文件编号  
        current_start += 1  # 增加文件编号  
    
    record_column = []  # 用于存储记录编号  
    group_size = records_count + 1  # 每组的大小  
    
    # 为每组生成记录编号  
    for group in range((row_num + group_size - 1) // group_size):  
        record_column.append(None)  # 每组的第一个元素是空  
        
        # 生成后续的元素，从 1 开始  
        for i in range(1, records_count + 1):  # 每组有 y 个计数元素  
            element_number = group * group_size + i  
            if element_number < row_num:  # 确保不超过总元素数量  
                record_column.append(str(i))  # 从 1 开始计数  
            else:  
                break  # 如果超出总元素数量，停止添加  
    
    # 将生成的文件编号和记录编号添加到新 DataFrame 中  
    new_df[file_column_name] = file_column  
    new_df[record_column_name] = record_column  
    
    return new_df  # 返回新的 DataFrame

def copy_columns_to_another(file_path, columns_mapping):
    """
    在 Excel 表中将指定的源列数据复制到对应的目标列，并更新目标列的值，而不是覆盖整个表格。

    参数:
    - file_path: Excel 文件的路径。
    - columns_mapping: 一个字典，其中键是源列的名称，值是目标列的名称。

    功能描述:
    1. 读取指定路径的 Excel 文件到 DataFrame。
    2. 遍历字典中的每个键值对，其中键是源列名，值是目标列名。
    3. 对于每个键值对，检查源列是否存在于 DataFrame 中，如果不存在则打印错误信息。
    4. 检查目标列是否已经存在于 DataFrame 中，如果不存在则创建一个新的目标列，并初始化所有值为 None。
    5. 将源列的数据复制到目标列，更新目标列的值。
    6. 将更新后的 DataFrame 保存回 Excel 文件，确保不包含行索引。
    7. 打印成功信息，告知用户列数据复制操作已完成。
    """
    # 读取Excel文件到DataFrame
    df = pd.read_excel(file_path)
    
    # 遍历字典中的每个键值对
    for source_column, target_column in columns_mapping.items():
        # 检查源列是否存在于DataFrame中
        if source_column not in df.columns:
            print(f"源列 '{source_column}' 不存在于Excel文件中。")
            continue  # 跳过当前循环，继续下一个
        
        # 检查目标列是否已经存在于DataFrame中
        if target_column not in df.columns:
            df[target_column] = None
        
        # 将源列的数据复制到目标列
        df[target_column] = df[source_column]
    
    # 将更新后的DataFrame保存回Excel文件
    df.to_excel(file_path, index=False)
    
    # 打印成功信息，告知用户列数据复制操作已完成
    print("列数据复制成功。")

def copy_strings_to_excel_columns(file_path, columns_values):
    """
    将指定的字符串复制到Excel表的多个指定列。

    参数:
    - file_path: Excel 文件的路径。
    - columns_values: 一个字典，其中键是目标列的名称，值是要复制的字符串值。

    功能描述:
    1. 读取指定路径的 Excel 文件到 DataFrame。
    2. 遍历字典中的每个键值对，其中键是目标列名，值是要复制的字符串。
    3. 对于每个键值对，检查目标列是否已经存在于 DataFrame 中，如果不存在则创建一个新的目标列，并初始化所有值为指定的字符串。
    4. 如果目标列已存在，则直接将指定的字符串值填充到整列。
    5. 将更新后的 DataFrame 保存回 Excel 文件，确保不包含行索引。
    6. 打印成功信息，告知用户字符串复制操作已完成。
    """
    # 读取Excel文件到DataFrame
    df = pd.read_excel(file_path)
    
    # 遍历字典中的每个键值对
    for target_column, value in columns_values.items():
        # 检查目标列是否已经存在于DataFrame中
        if target_column not in df.columns:
            # 如果目标列不存在，则创建一个新的目标列，并初始化所有值为指定的字符串
            df[target_column] = value
        else:
            # 如果目标列已存在，则直接将指定的字符串值填充到整列
            df[target_column] = value
    
    # 将更新后的DataFrame保存回Excel文件
    df.to_excel(file_path, index=False)
    
    # 打印成功信息，告知用户字符串复制操作已完成
    print("字符串复制成功。")    

def join_columns(df, column_list, target_column):
    """
    此函数用于处理给定数据帧，将指定列按照特定规则连接后放入指定的目标列中。

    参数：
    df：需要处理的 DataFrame 对象。
    column_list：包含列名的列表，用于指定要连接的列。
    target_column：目标列名，连接后的结果将放入此列。

    返回：
    处理后的 DataFrame 对象。
    """
    for index, row in df.iterrows():
        # 如果列表中只有一列，则直接将该列的值放入目标列，不进行连接操作
        if len(column_list) == 1:
            df.at[index, target_column] = row[column_list[0]]
        else:
            # 获取除最后一列外的列名列表
            columns_except_last = column_list[:-1]
            # 连接除最后一列外的列的值
            joined_string = '-'.join(str(row[col]) for col in columns_except_last if pd.notna(row[col]))
            # 如果最后一列存在
            if column_list[-1] in df.columns and pd.notna(row[column_list[-1]]):
                last_column_value = str(int(row[column_list[-1]]))
                # 使用点号连接最后一列的值
                joined_string += '.' + last_column_value
            df.at[index, target_column] = joined_string
    return df

# 检查文件夹是否存在，不存在就创建
def check_create_folder(folder_path, create_if_not_exists=True):

    desktop_path = os.path.expanduser("~/Desktop")
    
    # 检查文件夹路径是否为桌面
    if folder_path == desktop_path:
        print("文件夹路径为桌面:", folder_path)
        raise ValueError(f"文件夹 '{folder_path}' 是桌面，长此以往会让桌面一团糟。")
    
    # 检查文件夹是否存在
    if os.path.exists(folder_path):
        print("文件夹存在:", folder_path)
        return folder_path
    
    if create_if_not_exists:
        try:
            os.makedirs(folder_path)
            print("文件夹已创建:", folder_path)
            return folder_path
        except OSError as e:
            print("创建文件夹时出错:", e)
            return None
    else:
        print("文件夹不存在:", folder_path)
        return False

def list_files(directory, flag, exclude=None, output_excel=None, cover=False, parent=None):  
    """  
    列出指定目录下的所有文件和文件夹，可选择排除特定文件，并将结果以新列的形式插入到给定的Excel表的最前面。  

    参数:  
    - directory: 要列出文件的目录路径。  
    - flag: 如果为True，包括第一级子目录中的文件和文件夹；如果为False，仅包括当前目录下的文件和文件夹。  
    - exclude: 一个列表，包含要排除的文件或文件夹名称。  
    - output_excel: 给定的Excel文件路径，将结果插入到这个文件中。  
    - cover: 如果为True，覆盖现有数据；如果为False，在原有数据后面添加。  
    - parent: 用于显示 QMessageBox 的 QWidget 的引用。  
    """  
    # 确保目录存在  
    if not os.path.isdir(directory):  
        # QMessageBox.critical(parent, "错误", "指定的目录不存在。")  
        raise ValueError("The specified directory does not exist.")  

    # 如果没有提供排除列表，则创建一个空列表  
    if exclude is None:  
        exclude = []  

    # 准备数据存储列表  
    data = []  

    # 获取目录中的文件和文件夹列表  
    with os.scandir(directory) as entries:  
        for entry in entries:  
            if entry.name not in exclude:  
                # 去除扩展名并转换为字符串  
                file_name_without_extension = str(os.path.splitext(entry.name)[0])  
                if entry.is_file():  
                    # 获取文件的绝对路径  
                    file_path = entry.path  
                    print(file_path)
                    # 获取文件扩展名并转换为字符串  
                    extension = str(os.path.splitext(entry.name)[1])  
                    # 添加文件名称、路径、扩展名和超链接  
                    data.append([file_name_without_extension, file_path, extension, f'=HYPERLINK("{file_path}", "Open")'])  
                elif entry.is_dir() and flag:  
                    # 如果flag为True，包括第一级子目录的名称和路径  
                    folder_path = entry.path  
                    # 将文件夹名转换为字符串  
                    folder_name = str(entry.name)  
                    data.append([folder_name, folder_path, None, f'=HYPERLINK("{folder_path}", "Open")'])  

    # 创建DataFrame  
    df_files = pd.DataFrame(data, columns=['fileName', 'filePath', 'fileExtension', 'filePath_Link'])  

    if output_excel:  
        # 读取给定的Excel文件  
        if not os.path.isfile(output_excel):  
            # QMessageBox.critical(parent, "错误", "指定的Excel文件不存在。")  
            raise ValueError("The specified Excel file does not exist.")  
        df_excel = pd.read_excel(output_excel)  

        # 根据 cover 参数决定是覆盖还是添加数据  
        if cover:  
            # 仅覆盖 Name, Path, Extension, Link 列  
            df_excel[['fileName', 'filePath', 'fileExtension', 'filePath_Link']] = df_files[['fileName', 'filePath', 'fileExtension', 'filePath_Link']]  
        else:  
            # 在原有数据后面添加新数据  
            df_excel = pd.concat([df_excel, df_files], ignore_index=True)  

            # 确保 Link 列包含所有提取的链接  
            df_excel['filePath_Link'] = df_excel['filePath_Link'].combine_first(df_files['filePath_Link'])  

        # 确保列顺序为 Name, Path, Extension, Link  
        df_excel = df_excel[['fileName', 'filePath', 'fileExtension', 'filePath_Link'] + [col for col in df_excel.columns if col not in ['fileName', 'filePath', 'fileExtension', 'filePath_Link']]]  

        # 保存更新后的DataFrame到Excel文件  
        df_excel.to_excel(output_excel, index=False)  
        return output_excel
        # QMessageBox.information(parent, "完成", f"更新后的Excel文件已保存到 {output_excel}")  
    else:  
        # 保存到Excel文件  
        excel_path = os.path.join(directory, 'file_list.xlsx')  
        df_files.to_excel(excel_path, index=False)  
        return excel_path
        # QMessageBox.information(parent, "完成", f"文件列表已保存到 {excel_path}")

def list_files_recuision(directory, flag, exclude=None, output_directory=None):
    """
    列出指定目录下的所有文件和文件夹，可选择排除特定文件，并将结果保存到Excel文件中。
    此函数为递归收集信息
    参数:
    - directory: 要列出文件的目录路径。
    - flag: 如果为True，包括子目录中的文件；如果为False，仅包括当前目录下的文件。
    - exclude: 一个列表，包含要排除的文件或文件夹名称。
    - output_directory: 保存结果Excel文件的目录路径。如果未提供，则默认为directory。
    """
    # 确保目录存在
    if not os.path.isdir(directory):
        raise ValueError("The specified directory does not exist.")

    # 如果没有提供排除列表，则创建一个空列表
    if exclude is None:
        exclude = []

    # 确保输出目录存在，如果未提供，则使用输入目录
    if output_directory is None:
        output_directory = directory
    if not os.path.isdir(output_directory):
        raise ValueError("The specified output directory does not exist.")

    # 准备数据存储列表
    data = []

    # 遍历目录
    for root, dirs, files in os.walk(directory):
        for name in files:
            # 检查文件是否在排除列表中
            if name not in exclude:
                # 获取文件的绝对路径
                file_path = os.path.join(root, name)
                # 获取文件扩展名
                extension = os.path.splitext(name)[1]
                # 去除文件名中的扩展名
                file_name_without_extension = os.path.splitext(name)[0]
                data.append([file_name_without_extension, file_path, extension])
        
        if not flag:
            # 如果flag为False，只获取当前目录下的文件，不进入子目录
            break

        if flag:
            # 如果flag为True，获取所有一级子文件夹的名称和绝对路径
            for name in dirs:
                # 检查文件夹是否在排除列表中
                if name not in exclude:
                    folder_path = os.path.join(root, name)
                    data.append([name, folder_path, None])  # 文件夹没有扩展名

    # 创建DataFrame
    df = pd.DataFrame(data, columns=['Name', 'Path', 'Extension'])

    # 保存到Excel文件
    excel_path = os.path.join(output_directory, 'file_list.xlsx')
    df.to_excel(excel_path, index=False)

    print(f"File list saved to {excel_path}")

def rename_file_system_paths(df, original_path_col, target_path_col):
    """
    将 df 中指定的 original_path 列中的路径重命名为 target_path 列中的路径，
    实际在文件系统中进行重命名，并检查原始和目标路径列中的重复项。

    :param df: 包含文件路径的 DataFrame
    :param original_path_col: 表示原始路径的列名
    :param target_path_col: 表示目标路径的列名
    :return: 带有操作结果的 DataFrame，包括成功或失败的信息
    """
    print('ok')
    results = []

    # 检查 original_path 和 target_path 列中的重复项
    if df[original_path_col].duplicated().any():
        return pd.DataFrame({'status': ['error'], 'message': ['Original path column has duplicates.']})

    if df[target_path_col].duplicated().any():
        return pd.DataFrame({'status': ['error'], 'message': ['Target path column has duplicates.']})

    for index, row in df.iterrows():
        try:

            original_path = row[original_path_col]
            target_path = row[target_path_col]

            # 检查 original_path 是否存在，如果不存在就忽略该行处理
            if not os.path.exists(original_path):
                continue

            # 检查目标路径是否已经存在
            if os.path.exists(target_path):
                results.append({'index': index, 'status': 'error', 'message': f'Target path already exists: {target_path}'})
                continue

            # 执行重命名操作
            os.rename(original_path, target_path)
            results.append({'index': index, 'status': 'success', 'message': f'Renamed {original_path} to {target_path}'})
        except Exception as e:
            # 更详细的错误信息记录
            error_message = f'Error at index {index}: {str(e)}'
            results.append({'index': index, 'status': 'error', 'message': error_message})

    # 将结果添加到 DataFrame 中
    result_df = pd.DataFrame(results)
    return result_df

def rename_files(excel_path, new_name_column, directory, parent=None):  
    """  
    根据 Excel 表中的信息重命名指定目录下的文件或文件夹。  

    参数:  
    - excel_path: Excel 表格的路径，包含文件的当前名称、扩展名、完整路径和新名称。  
    - new_name_column: Excel 表格中包含新名称的列的列名。  
    - directory: 要重命名文件或文件夹的目录路径。  
    - parent: 用于显示 QMessageBox 的 QWidget 的引用。  
    """  
    success = False  # 初始化标志变量  

    # 确保目录存在  
    if not os.path.isdir(directory):  
        # if parent:  

        #     QMessageBox.critical(parent, "错误", "The specified directory does not exist.")  
        # else:  
        #     QMessageBox.critical(None, "错误", "The specified directory does not exist.")  
        raise ValueError("The specified directory does not exist.")  

    # 读取 Excel 表格  
    try:  
        df = pd.read_excel(excel_path)  
    except Exception as e:  
        # if parent:  
        #     QMessageBox.critical(parent, "错误", f"Failed to read Excel file: {e}")  
        # else:  
        #     QMessageBox.critical(None, "错误", f"Failed to read Excel file: {e}")  
        raise ValueError(f"Failed to read Excel file: {e}")  

    # 检查必要的列是否存在  
    if 'fileName' not in df.columns or 'fileExtension' not in df.columns or 'filePath' not in df.columns or new_name_column not in df.columns:  
        # if parent:  
        #     QMessageBox.critical(parent, "错误", f"Excel file must contain 'fileName', 'fileExtension', 'filePath', and {new_name_column} columns.")  
        # else:  
        #     QMessageBox.critical(None, "错误", f"Excel file must contain 'fileName', 'fileExtension', 'filePath', and {new_name_column} columns.")  
        raise ValueError(f"Excel file must contain 'Name', 'Extension', 'Path', and {new_name_column} columns.")  

    # 检查 Name 列的数量是否与 new_name_column 列的数量一致  
    if df['fileName'].count() != df[new_name_column].count():  
        error_message = "The number of 'Name' entries does not match the number of entries in the new name column."  
        # if parent:  
        #     QMessageBox.critical(parent, "错误", error_message)  
        # else:  
        #     QMessageBox.critical(None, "错误", error_message)  
        return  # 终止程序  

    # 检查 new_name_column 列中的名称是否唯一且不为空  
    if df[new_name_column].isnull().any():  
        error_message = f"The column {new_name_column} contains null values."  
        # if parent:  
        #     QMessageBox.critical(parent, "错误", error_message)  
        # else:  
        #     QMessageBox.critical(None, "错误", error_message)  
        return  # 终止程序  

    # 使用 Path 列检查新名称的唯一性  
    if len(df['filePath'].unique()) != len(df['filePath']):  
        error_message = f"The column 'Path' contains duplicate paths."  
        # if parent:  
        #     QMessageBox.critical(parent, "错误", error_message)  
        # else:  
        #     QMessageBox.critical(None, "错误", error_message)  
        return  # 终止程序  

    # 检查同一目录下 new_name_column + extension 的唯一性  
    df['NewFullName'] = df[new_name_column] + df['fileExtension'].fillna("")  # 创建新名称列  
    existing_files = {f for f in os.listdir(directory)}  # 获取目录下现有文件  
    if any(f"{row[new_name_column]}{row['fileExtension']}" in existing_files for _, row in df.iterrows()):  
        error_message = f"The combination of new names and extensions contains duplicates in the directory."  
        if parent:  
            QMessageBox.critical(parent, "错误", error_message)  
        else:  
            QMessageBox.critical(None, "错误", error_message)  
        return  # 终止程序  

    # 遍历 Excel 表中的每一行  
    for index, row in df.iterrows():  
        # 从 Excel 表中获取文件或文件夹的名称和扩展名，并确保它们是字符串类型  
        name = str(row['fileName'])  
        extension = str(row['fileExtension']).strip() if pd.notnull(row['fileExtension']) else ""  # 处理 NaN 和空格  
        new_name = str(row[new_name_column])  

        # 构建原始和新的文件或文件夹路径  
        if extension:  # 如果扩展名不为空，表示它是一个文件  
            original_file_name = f"{name}{extension}"  
            original_path = os.path.join(directory, original_file_name)  
            new_file_name = f"{new_name}{extension}"  
            new_path = os.path.join(directory, new_file_name)  
        else:  # 如果扩展名为空，表示它是一个文件夹  
            original_path = os.path.join(directory, name)  
            new_path = os.path.join(directory, new_name)  

        # 检查原始文件或文件夹是否存在  
        if os.path.exists(original_path):  
            # 重命名文件或文件夹  
            os.rename(original_path, new_path)  
            success = True  # 标记至少有一个文件被成功重命名  
            print(f"Renamed '{original_path}' to '{new_path}'")  
        else:  
            print(f"File or folder '{original_path}' does not exist. Skipping.")  

    # 如果有文件被成功重命名，弹出一次成功对话框  
    return True
    # if success and parent:  
    #     QMessageBox.information(parent, "成功", "Files have been successfully renamed.")  
    # elif success:  
    #     QMessageBox.information(None, "成功", "Files have been successfully renamed.")

# 检查指定文件是否存在
def check_file_existence(file_path:str)->bool:
    if os.path.exists(file_path):
        print(f"文件存在： {file_path} ")
        return True
    else:
        print(f"文件不存在： {file_path} ")
        return False
        
def check_files(df, directory):  
    """  
    检查 Excel 表中的文件名（包括扩展名）是否与给定目录中的一级子目录和文件名一一对应。  

    参数:    - df: dataframe数据，包含文件名和扩展名。  
    - directory: 要检查的目录路径。  

    返回:  
    - bool: 如果文件名完全对应则返回 True，否则返回 False。  
    - dict: 目录中存在但 Excel 表中不存在的文件名（包括扩展名）及其绝对路径。  
    - dict: Excel 表中存在但目录中不存在的文件名（包括扩展名）及其绝对路径。  
    """  
    # 确保目录存在  
    if not os.path.isdir(directory):  
        raise ValueError(f"The specified directory does not exist: {directory}")  

    # 检查 Excel 表中是否包含 'Name' 和 'Extension' 列  
    if 'Name' not in df.columns or 'Extension' not in df.columns:  
        raise ValueError("Excel file must contain 'Name' and 'Extension' columns.")  

    # 获取目录中的一级子目录和文件名  
    directory_items = set(os.listdir(directory))  

    # 获取 Excel 表中的文件名（包括扩展名）  
    excel_items = set()  
    for index, row in df.iterrows():  
        name = str(row['Name'])  
        extension = str(row['Extension']).strip() if pd.notnull(row['Extension']) else ""  # 处理 NaN 和空格  
        full_name = f"{name}{extension}" if extension else name  
        excel_items.add(full_name)  

    # 找出目录中存在但 Excel 表中不存在的文件名和文件夹名  
    only_in_directory = {os.path.join(directory, item): item for item in directory_items - excel_items}  

    # 找出 Excel 表中存在但目录中不存在的文件名  
    only_in_excel = {os.path.join(directory, item): item for item in excel_items - directory_items}  

    # 检查是否完全对应  
    if not only_in_directory and not only_in_excel:  
        return True, {}, {}  
    else:  
        return False, only_in_directory, only_in_excel  

def is_archive_code(str):
    """
    判断字符串是否符合2023-JX14.13-1或2023-JX14.13-0001或2023-JX14.13-1.1或2023-JX14.13-1.0001或2023-JX14.13-0001.1或2023-JX14.13-0001.0001或2023-JX14.GZ-0001.0001
    参数:
    - str: 要检查的字符串。

    返回:
    - bool: 如果字符串符合格式则返回 True，否则返回 False。
    """
    pattern = r'^((19|20)\d{2})-[a-zA-Z]{2}(\d{2})\.([A-Za-z]{2}(\d)?|\d{2})(-\d{1,4})(\.\d{1,4})?$'
    return re.match(pattern, str) is not None

def is_class_code(str):
    """
    判断字符串是否符合JX14.13或JX14.GZ或JX14.GZ0
    参数:
    - str: 要检查的字符串。

    返回:
    - bool: 如果字符串符合格式则返回 True，否则返回 False。
    """
    pattern = r'^[a-zA-Z]{2}(\d{2})\.([A-Za-z]{2}(\d)?|\d{2})$'
    return re.match(pattern, str) is not None

def is_year(str):
    """
    判断字符串是否符合年度形式
    参数:
    - str: 要检查的字符串。

    返回:
    - bool: 如果字符串符合格式则返回 True，否则返回 False。
    """
    pattern = r'^((19|20)\d{2})$'
    return re.match(pattern, str) is not None

def split_archive_code(archive_code):
    """
    将符合特定格式的字符串拆分为数组。

    参数:
    - archive_code: 要拆分的字符串。

    返回:
    - list: 拆分后的数组，如果字符串不符合格式则返回 None。
    """
    pattern = r'^((19|20)\d{2})-[a-zA-Z]{2}(\d{2})\.([A-Za-z]{2}(\d)?|\d{2})(-\d{1,4})(\.\d{1,4})?$'
    match = re.match(pattern, archive_code)
    if match:
        # 拆分字符串
        return archive_code.split('-')
    else:
        return None
    
def extract_da_type(text):
    """
    从给定的字符串中提取前两个大写字母。

    参数:
    - text: 要提取的字符串。

    返回:
    - str: 提取的前两个大写字母，如果不符合模式则返回 None。
    """
    pattern = r'^([a-zA-Z]{2})\d{2}\.([A-Za-z]{2}(\d)?|\d{2})$'
    match = re.match(pattern, text)
    if match:
        return match.group(1)
    else:
        return None
    
def split_and_normalize_numbers(s):
    """
    将字符串按点分割成数组，并去除数字元素的无意义前导零。

    参数:
    - s: 要分割和处理的字符串。

    返回:
    - list: 处理后的数字数组。
    """
    if '.' not in s:
        normalized_parts = [str(int(s)) if s.isdigit() else s]  # 去除前导零
    else:
        # 使用正则表达式分割字符串
        parts = s.split('.')
        
        # 处理每个部分，去除无意义的前导零
        normalized_parts = []
        for part in parts:
            # 检查是否为数字
            if part.isdigit():
                # 转换为整数，然后转回字符串，以去除前导零
                normalized_part = str(int(part))
            else:
                # 非数字部分保持不变
                normalized_part = part
            normalized_parts.append(normalized_part)
    
    return normalized_parts

def source_to_target_path_dict(base_dir,excel_path,column_name):
    """
    从给定的Excel文件路径读取数据，只包含指定的字段列表，并返回DataFrame。

    参数:
    - base_dir: 基础目录。
    - excel_path: Excel文件的路径。
    - column_name: 字段名列表。
    返回:
    - 一个字典。
    """
    final_path_dict={}
    # 尝试读取Excel文件
    try:
        # 使用pandas读取Excel文件
        df = pd.read_excel(excel_path, usecols=column_name)
        # 检查DataFrame是否已正确加载
        if df is not None:
            # 可以在这里对df的每一行进行进一步处理
            for index, row in df.iterrows():
                if is_archive_record_code(row['档号(必填)']): 
                    # 检查 'Extension' 是否为 NaN，如果是，则使用空字符串，否则转换为字符串
                    extension_str = '' if pd.isna(row['Extension']) else str(row['Extension'])
                    target_path=os.path.join(record_code_to_path(base_dir, str(row['档号(必填)'])),str(int(row['件号']))+' '+str(row['题名(必填)'])+extension_str)
                    source_path=str(row['Path'])
                    final_path_dict[source_path] =target_path
    except Exception as e:
        print(f"读取Excel文件时发生错误: {e}")
        return None
    
    return final_path_dict

def record_code_to_path(base_dir, code):  
    """  
    根据给定字符串创建目录结构。  

    参数:  
    - base_dir: 基础目录。  
    - code: 档号字符串，根据这些字符串创建子目录。 
    """  
    final_paths = []  # 初始化一个空列表来存储所有的 final_path  
    if is_archive_record_code(code):  
        parts = split_archive_code(code)  
        if parts:  
            year = parts[0]  # 年份部分  
            class_code = parts[1]  
            da_type = extract_da_type(class_code)  # 代码类型（包括字母和数字）  
            final_path = [base_dir, da_type, year, class_code]  # 构建路径初步  

            num = split_and_normalize_numbers(parts[2])  
            if len(num) == 2:  
                file_num, record_num = num  
            else:  
                file_num = num[0] if num else None  
                record_num = None  
            # 完整的目录路径  
            final_path.append(file_num)   
            # 使用 os.path.join 连接所有部分  
            final_path = os.path.join(*filter(None, final_path))  # 过滤掉 None 项  
            # 如果目录不存在，创建目录  
            if not os.path.exists(final_path):  
                os.makedirs(final_path)  
                print(f"Created directory: {final_path}")  
            else:  
                print(f"Directory already exists: {final_path}")  
    else:  
        print(f"Invalid archive code format: {code}")  
    return final_path  # 返回所有创建的路径列表

def delete_files_dirs(file_dir_list):
    """
    删除给定列表中的文件和文件夹。

    :param file_dir_list: 包含文件和文件夹路径的列表
    """
    for item_path in file_dir_list:
        try:
            if os.path.isfile(item_path):
                os.remove(item_path)
                print(f"Deleted file: {item_path}")
            elif os.path.isdir(item_path):
                os.rmdir(item_path)
                print(f"Deleted directory: {item_path}")
        except Exception as e:
            print(f"Error deleting {item_path}: {str(e)}")

def copy_files_or_directories(source_target_mapping):  
    """  
    将源目录或文件复制到目标路径。  

    参数:  
    - source_target_mapping: 一个字典，键是源路径（source_path），值是目标路径（target_path）。  

    功能描述:  
    1. 遍历字典中的每个键值对，键是源路径，值是目标路径。  
    2. 检查源路径是否存在，如果不存在则打印错误信息并跳过当前项。  
    3. 检查源路径是文件还是目录。  
    4. 如果是文件，则直接复制到目标路径。  
    5. 如果是目录，则递归地复制目录和其下所有子目录和文件到目标路径。  
    6. 进行必要的容错处理，如目标路径已存在时的处理。  
    """  
    for source_path, target_path in source_target_mapping.items():  
        # 检查源路径是否存在  
        if not os.path.exists(source_path):  
            print(f"源路径 '{source_path}' 不存在。")  
            continue  
        
        # 检查源路径是文件还是目录  
        if os.path.isfile(source_path):  
            # 源路径是文件，直接复制文件  
            shutil.copy(source_path, target_path)  
            print(f"文件 '{source_path}' 已复制到 '{target_path}'。")  
        elif os.path.isdir(source_path):  
            # 源路径是目录，复制目录及其所有子目录和文件  
            if os.path.exists(target_path):  
                print(f"目标路径 '{target_path}' 已存在，将直接覆盖该目录。")  
            else:  
                os.makedirs(target_path, exist_ok=True)  # 创建目标目录（如果不存在）  
            shutil.copytree(source_path, target_path, dirs_exist_ok=True)  
            print(f"目录 '{source_path}' 及其内容已复制到 '{target_path}'。")  
        else:  
            print(f"路径 '{source_path}' 不是文件或目录。") 
            return False
    return True

def is_archive_file_code(str):
    """
    判断字符串是否符合案卷档号格式2023-JX14.13-0001或2023-JX14.13-1或2023-JX14.13-0001或2023-JX14.GZ-0001
    参数:
    - str: 要检查的字符串。

    返回:
    - bool: 如果字符串符合格式则返回 True，否则返回 False。
    """
    pattern = r'^((19|20)\d{2})-[a-zA-Z]{2}(\d{2})\.([A-Za-z]{2}(\d)?|\d{2})(-\d{1,4})?$'
    return re.match(pattern, str) is not None

def is_archive_record_code(str):
    """
    判断字符串是否符合件档号格式或2023-JX14.13-1.1或2023-JX14.13-1.0001或2023-JX14.13-0001.1或2023-JX14.13-0001.0001或2023-JX14.GZ-0001.0001
    参数:
    - str: 要检查的字符串。

    返回:
    - bool: 如果字符串符合格式则返回 True，否则返回 False。
    """
    pattern = r'^((19|20)\d{2})-[a-zA-Z]{2}(\d{2})\.([A-Za-z]{2}(\d)?|\d{2})(-\d{1,4}\.\d{1,4})$'
    return re.match(pattern, str) is not None

def extract_str_from_pattern(strings_list, pattern):
    """
    从给定列表中提取符合正则表达式模式的字符串。

    参数:
    - strings_list: 原始字符串列表。
    - pattern: 用于匹配和提取的正则表达式模式。

    返回:
    - 一个新列表，包含修改后的字符串。
    """
    # 定义一个函数，用于替换匹配到的字符串
    def replace_match(match):
        # 提取匹配到的字符串，去除后面的数字
        return match.group(1) + match.group(2) + match.group(3).rstrip('.0')

    # 编译正则表达式模式
    regex = re.compile(pattern)

    # 使用列表推导式和正则表达式的 sub 方法提取和替换字符串
    extracted_list = [regex.sub(replace_match, item) for item in strings_list]

    return extracted_list

def create_directory_structure(base_dir, codes,flag=True):  
    """  
    根据给定字符串创建目录结构。  

    参数:  
    - base_dir: 基础目录。  
    - codes: 档号字符串列表，根据这些字符串创建子目录。
    - flag: 如果flag为True则建立到卷目录，否则只建立档案类型、年、分类号目录。 
    """  
    for code in codes:  
        if is_archive_code(code):  
            parts = split_archive_code(code)  
            if parts:  
                year = parts[0]  # 年份部分  
                class_code=parts[1]
                da_type = extract_da_type(class_code)  # 代码类型（包括字母和数字）
                final_path = [base_dir, da_type, year, class_code]  # 构建路径初步  

                if flag==True:
                    num=split_and_normalize_numbers(parts[2])
                    if len(num) == 2:
                        file_num, record_num = num
                    else:
                        file_num = num[0] if num else None
                        record_num = None
                    # 完整的目录路径  
                    final_path.append(file_num) 
                    # if record_num:  # 如果有记录号，加入路径  
                    #     final_path.append(record_num)  

                # 使用 os.path.join 连接所有部分  
                final_path = os.path.join(*filter(None, final_path))  # 过滤掉 None 项  

                # 如果目录不存在，创建目录  
                if not os.path.exists(final_path):  
                    os.makedirs(final_path)  
                    print(f"Created directory: {final_path}")  
                else:  
                    print(f"Directory already exists: {final_path}")  

        else:  
            print(f"Invalid archive code format: {code}")   

def highlight_missing_values(df1, df2, columns):
    """
    该函数用于比较两个 DataFrame 中指定列的值，找出彼此缺失的值，并将这些缺失值在原 DataFrame 中高亮显示。

    参数:
    df1 (pandas.DataFrame): 第一个参与比较的 DataFrame。
    df2 (pandas.DataFrame): 第二个参与比较的 DataFrame。
    columns (list): 要进行比较的列名列表。

    返回:
    tuple: 包含以下元素的元组
        - result_bool (bool): 表示两个 DataFrame 中彼此缺失值的集合是否相同。
        - styled_df1 (pandas.io.formats.style.Styler): 对 df1 中指定列的缺失值进行高亮显示后的样式化 DataFrame。
        - styled_df2 (pandas.io.formats.style.Styler): 对 df2 中指定列的缺失值进行高亮显示后的样式化 DataFrame。
        - values_in_df1_not_in_df2_all_style (pandas.io.formats.style.Styler): 包含在 df1 中存在但在 df2 中不存在的值的样式化 DataFrame。
        - values_in_df2_not_in_df1_all_style (pandas.io.formats.style.Styler): 包含在 df2 中存在但在 df1 中不存在的值的样式化 DataFrame。
    """
    # 将所有数据转换为字符串类型，以便统一比较
    df1 = df1.astype(str)
    df2 = df2.astype(str)

    # 创建用于存储缺失值的集合，初始化为空 DataFrame
    values_in_df1_not_in_df2_all = pd.DataFrame()
    values_in_df2_not_in_df1_all = pd.DataFrame()

    # 遍历指定的列
    for column in columns:
        # 在 df1 中存在但在 df2 中不存在的值，并重置索引
        values_in_df1_not_in_df2 = df1[column][~df1[column].isin(df2[column])].reset_index(drop=True)
        # 在 df2 中存在但在 df1 中不存在的值，并重置索引
        values_in_df2_not_in_df1 = df2[column][~df2[column].isin(df1[column])].reset_index(drop=True)

        # 创建新的 DataFrame 明确标识列名，并将其添加到总的缺失值集合中
        if not values_in_df1_not_in_df2.empty:
            temp_df1 = pd.DataFrame({column: values_in_df1_not_in_df2})
            values_in_df1_not_in_df2_all = pd.concat([values_in_df1_not_in_df2_all, temp_df1], ignore_index=True)

        if not values_in_df2_not_in_df1.empty:
            temp_df2 = pd.DataFrame({column: values_in_df2_not_in_df1})
            values_in_df2_not_in_df1_all = pd.concat([values_in_df2_not_in_df1_all, temp_df2], ignore_index=True)

    # 比较两个 DataFrame 是否相同，即彼此缺失值的集合是否一致
    result_bool = values_in_df1_not_in_df2_all.equals(values_in_df2_not_in_df1_all)

    # 为非 NaN 部分填充黄色的样式函数
    def highlight_non_nan(s):
        """
        该函数用于为 Series 中的非 NaN 值设置黄色背景样式。

        参数:
        s (pandas.Series): 要应用样式的 Series。

        返回:
        list: 包含样式字符串的列表，用于设置每个元素的样式。
        """
        return ['background-color: yellow' if pd.notna(v) else '' for v in s]

    def highlight_df1(row):
        """
        该函数用于为 df1 的每一行设置样式，将指定列中的缺失值高亮显示为黄色。

        参数:
        row (pandas.Series): df1 中的一行数据。

        返回:
        list: 包含样式字符串的列表，用于设置该行每个元素的样式。
        """
        styles = []
        # 遍历 df1 的所有列
        for column in df1.columns:
            value = row[column]
            # 只在指定的列中检查
            if column in columns:
                if value in values_in_df1_not_in_df2_all[column].values:
                    styles.append('background-color: yellow')
                else:
                    styles.append('')
            else:
                styles.append('')  # 对于未指定的列，不应用任何样式
        return styles

    def highlight_df2(row):
        """
        该函数用于为 df2 的每一行设置样式，将指定列中的缺失值高亮显示为黄色。

        参数:
        row (pandas.Series): df2 中的一行数据。

        返回:
        list: 包含样式字符串的列表，用于设置该行每个元素的样式。
        """
        styles = []
        # 遍历 df2 的所有列
        for column in df2.columns:
            value = row[column]
            # 只在指定的列中检查
            if column in columns:
                if value in values_in_df2_not_in_df1_all[column].values:
                    styles.append('background-color: yellow')
                else:
                    styles.append('')
            else:
                styles.append('')  # 对于未指定的列，不应用任何样式
        return styles

    # 应用样式到 df1 和 df2
    styled_df1 = df1.style.apply(highlight_df1, axis=1)
    styled_df2 = df2.style.apply(highlight_df2, axis=1)

    # 应用样式到缺失值集合
    values_in_df1_not_in_df2_all_style = values_in_df1_not_in_df2_all.style.apply(highlight_non_nan, axis=0)
    values_in_df2_not_in_df1_all_style = values_in_df2_not_in_df1_all.style.apply(highlight_non_nan, axis=0)

    return result_bool, styled_df1, styled_df2, values_in_df1_not_in_df2_all_style, values_in_df2_not_in_df1_all_style

def compare_files_in_excel_and_directory(excel_path, compare_dir, result_dir, qmessage=True):
    """
    该函数用于比较 Excel 文件中的文件信息和指定目录下的文件，找出彼此独有的文件，
    并将目录中多出的文件移动到指定结果目录，同时保存相关信息到 Excel 文件。

    参数:
    excel_path (str): Excel 文件的路径。
    compare_dir (str): 要进行比较的目录路径。
    result_dir (str): 保存结果的目录路径。
    qmessage (bool, 可选): 是否显示消息框，默认为 True。

    返回:
    tuple: 包含以下元素的元组
        - remains_path (str): 存放目录中多出文件的目录路径。
        - only_in_excel_path (str): 存放只在 Excel 中出现的文件信息的 Excel 文件路径。
    """
    try:
        # 读取 Excel 文件
        df = pd.read_excel(excel_path, header=0)

        # 比较文件的函数，假设已经定义
        bool_result, only_in_directory, only_in_excel = check_files(df, compare_dir)

        # 获取当前日期和时间
        now = datetime.datetime.now()
        date_string = now.strftime("%Y-%m-%d %H%M%S")
        remains_dir = 'remains ' + date_string + ' ' + secrets.token_hex(4)

        # 设置路径
        remains_path = os.path.join(result_dir, remains_dir)
        remains_catalog_file_path = os.path.join(remains_path, 'remains_catalogue.xlsx')
        only_in_excel_path = os.path.join(remains_path, 'only_in_excel.xlsx')

        # 确保路径是合法的
        os.makedirs(remains_path, exist_ok=True)

        # 处理目录中的文件，仅保留文件名
        for key in only_in_directory.keys():
            if os.path.isdir(key):
                only_in_directory[key] = ""

        n = len(only_in_directory)
        source_path = list(only_in_directory.keys())
        move_to_dir = [remains_path for _ in range(n)]
        files_name = list(only_in_directory.values())

        # 使用 zip() 创建目标路径
        target_path = [os.path.join(x, y) for x, y in zip(move_to_dir, files_name)]
        path_mapping = dict(zip(source_path, target_path))

        # 移动文件或目录的函数，假设已经定义
        move_files_or_directories(path_mapping)

        # 将路径映射结果写入 Excel
        remains_df = pd.DataFrame(list(path_mapping.items()), columns=['source_path', 'target_path'])
        remains_df.to_excel(remains_catalog_file_path, index=False)

        # 将只在 Excel 中找到的文件存入 DataFrame 并导出到 Excel
        only_in_excel_df = pd.DataFrame(list(only_in_excel.items()), columns=['source_path', 'name'])
        only_in_excel_df.to_excel(only_in_excel_path, index=False)
        print(f"目录中多出的文件已保存在: {remains_path}; \n EXCEL 表中多出的文件已保存在{only_in_excel_path}")
        # if qmessage:
        #     QMessageBox.information(None, "信息", f"目录中多出的文件已保存在: {remains_path}; \n EXCEL表中多出的文件已保存在{only_in_excel_path}")
        return remains_path, only_in_excel_path
    except Exception as e:
        # if qmessage:
        #     QMessageBox.critical(None, "错误", f"{e}")
        raise RuntimeError(f"错误: {str(e)}")

def open_file(file_path):
    """
    该函数使用系统命令打开指定路径的文件。

    参数:
    file_path (str): 要打开的文件的路径。

    返回:
    无
    """
    # 使用系统命令打开文件
    os.system(f"start {file_path}")

def find_duplicate_values(df, columns):
    """
    检查 DataFrame 中指定列的每行数据是否有重复值。
    如果有重复值，返回两个字典，
    第一个字典键是重复的值，值是包含重复单元格自定义地址的列表；
    第二个字典键是重复的值，值是包含重复单元格 Excel 地址的列表；
    如果没有重复值，返回 False。

    参数:
    df (pandas.DataFrame): 要检查的 DataFrame。
    columns (list): 要检查的列名列表。

    返回:
    tuple: 包含两个字典或布尔值的元组，有重复值时为 (dict1, dict2)，无重复值时为 (False, False)
    """
    # 容错处理：检查 df 是否为有效的 DataFrame
    if not isinstance(df, pd.DataFrame):
        print("输入的 df 不是有效的 DataFrame 类型，请检查输入。")
        return (False, False)

    # 容错处理：检查 columns 是否为有效的列表
    if not isinstance(columns, list):
        print("输入的 columns 不是有效的列表类型，请检查输入。")
        return (False, False)

    # 容错处理：检查 columns 中的列名是否存在于 df 中
    non_existent_columns = [col for col in columns if col not in df.columns]
    if non_existent_columns:
        print(f"指定的列名 {non_existent_columns} 不存在于 DataFrame 中，请检查列名。")
        return (False, False)

    custom_address_duplicates = {}
    excel_address_duplicates = {}
    has_duplicates = False

    for col in columns:
        try:
            # 将列中的每个值转换为字符串类型
            df[col] = df[col].astype(str)
            # 去除字符串中的所有空格
            df[col] = df[col].str.replace(" ", "")

            # 统计每列中每个值的出现次数
            value_counts = df[col].value_counts()
            # 筛选出出现次数大于 1 的值（即重复值）
            duplicate_values = value_counts[value_counts > 1].index

            for value in duplicate_values:
                has_duplicates = True
                # 找到该重复值在 DataFrame 中的所有行索引
                rows = df[df[col] == value].index
                # 生成自定义单元格地址列表
                custom_cell_addresses = [f"{col}{row + 1}" for row in rows]  # 行值加 1
                # 获取列的索引
                col_index = df.columns.get_loc(col) + 1
                # 生成 Excel 风格的单元格地址列表
                excel_cell_addresses = [f"{get_column_letter(col_index)}{row + 2}" for row in rows]  # 行值加 1

                if value in custom_address_duplicates:
                    custom_address_duplicates[value].extend(custom_cell_addresses)
                else:
                    custom_address_duplicates[value] = custom_cell_addresses

                if value in excel_address_duplicates:
                    excel_address_duplicates[value].extend(excel_cell_addresses)
                else:
                    excel_address_duplicates[value] = excel_cell_addresses

        except Exception as e:
            print(f"处理列 {col} 时出现错误: {e}")

    if has_duplicates:
        print("发现重复值，以下是重复值及其对应的自定义单元格地址：")
        for key, value in custom_address_duplicates.items():
            print(f"重复值: {key}, 自定义单元格地址: {', '.join(value)}")

        print("以下是重复值及其对应的 Excel 单元格地址：")
        for key, value in excel_address_duplicates.items():
            print(f"重复值: {key}, Excel 单元格地址: {', '.join(value)}")

        return (custom_address_duplicates, excel_address_duplicates)
    else:
        print("指定列中未发现重复值。")
        return (False, False)