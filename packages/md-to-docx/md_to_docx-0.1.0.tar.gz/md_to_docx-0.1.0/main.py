from typing import Any
import httpx
from mcp.server.fastmcp import FastMCP
import os
import re
import subprocess
import tempfile
import uuid
from pathlib import Path
import argparse
import requests
import json
import base64
from io import BytesIO
from bs4 import BeautifulSoup

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# Initialize FastMCP server
mcp = FastMCP("md_to_docx")


def render_mermaid_to_image(mermaid_code, output_path=None):
    """
    Render Mermaid diagram to an image using multiple methods.
    Returns the path to the saved image.
    """
    # Create a temporary file to save the image if not provided
    if not output_path:
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f"mermaid_{uuid.uuid4()}.png")
    
    # Method 1: Try using Mermaid.ink API
    try:
        # Encode the Mermaid code for the URL
        encoded_data = {"code": mermaid_code}
        json_str = json.dumps(encoded_data)
        base64_str = base64.urlsafe_b64encode(json_str.encode('utf-8')).decode('utf-8')
        
        # Use the Mermaid.ink API
        api_url = f"https://mermaid.ink/img/{base64_str}"
        response = requests.get(api_url)
        
        if response.status_code == 200:
            # Save the image
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Error using Mermaid.ink API: {e}")
    
    # Method 2: Try using mermaid-cli if available
    try:
        # Check if mmdc (mermaid-cli) is installed
        subprocess.run(["mmdc", "--version"], capture_output=True, check=True)
        
        # Create a temporary file for the Mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        # Run mmdc to generate the image
        subprocess.run([
            "mmdc",
            "-i", temp_mmd_path,
            "-o", output_path,
            "-b", "transparent"
        ], check=True)
        
        # Clean up the temporary Mermaid file
        os.unlink(temp_mmd_path)
        
        return output_path
    except (subprocess.SubprocessError, FileNotFoundError) as e:
        print(f"Error using mermaid-cli: {e}")
    
    # Method 3: Try using the Kroki API
    try:
        payload = {
            "diagram_source": mermaid_code,
            "diagram_type": "mermaid",
            "output_format": "png"
        }
        
        response = requests.post("https://kroki.io/", json=payload)
        
        if response.status_code == 200:
            # Save the image
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Error using Kroki API: {e}")
    
    # All methods failed
    print("All rendering methods failed for Mermaid diagram")
    return None


def extract_mermaid_blocks(md_content):
    """Extract Mermaid code blocks from Markdown content."""
    # Pattern to match ```mermaid ... ``` blocks
    pattern = r'```mermaid\s+(.*?)\s+```'
    # Find all matches using re.DOTALL to match across multiple lines
    matches = re.findall(pattern, md_content, re.DOTALL)
    return matches


def html_to_docx(html_content, doc):
    """Convert HTML content to Word document elements."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process elements in order
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'table']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = f'Heading {element.name[1]}'
        
        elif element.name == 'p':
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            apply_style_to_paragraph(paragraph, element)
        
        elif element.name == 'blockquote':
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = 'Quote'
        
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Bullet')
        
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Number')
        
        elif element.name == 'table':
            process_table(doc, element)
    
    return doc


def apply_style_to_paragraph(paragraph, element):
    """Apply HTML styles to a Word paragraph based on the element."""
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        paragraph.style = f'Heading {level}'
    
    if element.name == 'strong' or element.find('strong'):
        for run in paragraph.runs:
            run.bold = True
    
    if element.name == 'em' or element.find('em'):
        for run in paragraph.runs:
            run.italic = True
    
    if element.name == 'u' or element.find('u'):
        for run in paragraph.runs:
            run.underline = True
    
    if element.name == 'code' or element.find('code'):
        for run in paragraph.runs:
            run.font.name = 'Courier New'
    
    if element.name == 'center' or element.get('align') == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_list_item(doc, li_element, list_style, level=0):
    """Process a list item and its children with proper indentation."""
    # Add the list item with proper style and level
    text = li_element.get_text(strip=True)
    paragraph = doc.add_paragraph(text)
    paragraph.style = list_style
    paragraph.paragraph_format.left_indent = Pt(18 * level)  # Indent based on nesting level
    
    # Process any nested lists
    nested_ul = li_element.find('ul')
    nested_ol = li_element.find('ol')
    
    if nested_ul:
        for nested_li in nested_ul.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Bullet', level + 1)
    
    if nested_ol:
        for nested_li in nested_ol.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Number', level + 1)

def process_table(doc, table_element):
    """Process a table element and convert it to a Word table."""
    # Find all rows in the table
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count the maximum number of cells in any row
    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        max_cols = max(max_cols, len(cells))
    
    if max_cols == 0:
        return
    
    # Create the table in the document
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    # Fill the table with data
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < max_cols:  # Ensure we don't exceed the table dimensions
                # Get cell text and apply basic formatting
                text = cell.get_text(strip=True)
                table.cell(i, j).text = text
                
                # Apply header formatting if it's a header cell
                if cell.name == 'th' or i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True



@mcp.tool()
async def md_to_docx(md_content: str, output_file: str = None):
    """Convert Markdown file to DOCX, rendering Mermaid diagrams as images.
    
    Args:
        md_content: Markdown content to convert
        output_file: Optional output file path, defaults to 'output.docx'
    """
    # Read the Markdown file
    # with open(md_file, 'r', encoding='utf-8') as f:
    #     md_content = f.read()
    
    # Extract Mermaid blocks
    mermaid_blocks = extract_mermaid_blocks(md_content)
    
    # Create a new Word document
    doc = Document()
    
    # Replace Mermaid blocks with placeholders and keep track of them
    placeholders = []
    for i, block in enumerate(mermaid_blocks):
        placeholder = f"MERMAID_DIAGRAM_{i}"
        placeholders.append(placeholder)
        md_content = md_content.replace(f"```mermaid\n{block}\n```", placeholder)
    
    # Convert Markdown to HTML with extensions
    html_content = markdown.markdown(
        md_content,
        extensions=[
            'markdown.extensions.extra',
            'markdown.extensions.codehilite',
            'markdown.extensions.tables',
            'markdown.extensions.toc'
        ]
    )
    
    # Split HTML by placeholders
    parts = []
    for part in re.split(f"({'|'.join(placeholders)})", html_content):
        if part in placeholders:
            # This is a placeholder, mark it for later replacement
            parts.append((True, placeholders.index(part)))
        else:
            # This is regular HTML content
            parts.append((False, part))
    
    # Process each part
    for is_placeholder, content in parts:
        if is_placeholder:
            # Render the Mermaid diagram
            mermaid_code = mermaid_blocks[content]
            img_path = render_mermaid_to_image(mermaid_code)
            
            if img_path:
                # Add the image to the document
                doc.add_picture(img_path, width=Inches(6))
                
                # Clean up the temporary image file
                try:
                    os.unlink(img_path)
                except:
                    pass
            else:
                # If rendering failed, add the Mermaid code as text
                doc.add_paragraph("Failed to render Mermaid diagram:", style='Intense Quote')
                code_para = doc.add_paragraph(mermaid_code)
                code_para.style = 'No Spacing'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        else:
            # Add regular content as paragraphs with proper formatting
            if content.strip():
                html_to_docx(content, doc)
    
    # Determine output file name if not provided
    if not output_file:
        output_file = 'output.docx'
    
    # Save the document
    doc.save(output_file)
    return output_file


if __name__ == "__main__":
    # Initialize and run the server
    mcp.run(transport='stdio')