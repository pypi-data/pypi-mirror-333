from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO
from bs4 import BeautifulSoup, NavigableString
import cssutils

class H2D:
    def __init__(self, font_name_or_path):
        self.document = None
        self.current_paragraph = None
        self.list_level = 0
        self.list_type = None
        self.list_counters = []
        self.font_name_or_path = font_name_or_path
        self.tag_handlers = {
            'h1': lambda e: self.handle_heading(e, 1),
            'h2': lambda e: self.handle_heading(e, 2),
            'h3': lambda e: self.handle_heading(e, 3),
            'h4': lambda e: self.handle_heading(e, 4),
            'h5': lambda e: self.handle_heading(e, 5),
            'h6': lambda e: self.handle_heading(e, 6),
            'p': self.handle_paragraph,
            'a': self.handle_link,
            'strong': self.handle_strong,
            'em': self.handle_em,
            'ul': self.handle_ul,
            'ol': self.handle_ol,
            'li': self.handle_li,
            'span': self.handle_span,
            'br': self.handle_br,
            'hr': self.handle_hr,
            'header': self.handle_header,
            'footer': self.handle_footer,
            'nav': self.handle_nav,
            'article': self.handle_article,
            'section': self.handle_section,
            'aside': self.handle_aside,
            'figure': self.handle_figure,
            'figcaption': self.handle_figcaption,
            'main': self.handle_main,
            'mark': self.handle_mark,
            'time': self.handle_time,
            'pre': self.handle_pre,
            'img': self.handle_img,
            'table': self.handle_table,
            'tr': self.handle_tr,
            'td': self.handle_td,
            'th': self.handle_th,
        }

    def convert(self, html_string):
        """将 HTML 字符串转换为 DOCX Document 对象"""
        self.document = Document()
        self.ensure_styles()
        self.apply_global_font()
        self.list_level = 0
        self.list_type = None
        self.list_counters = []
        self.current_paragraph = None
        soup = BeautifulSoup(html_string, 'html.parser')
        for element in soup.children:
            self.handle_element(element)
        return self.document

    def ensure_styles(self):
        """确保文档中存在所需的样式"""
        styles = self.document.styles
        if 'Code' not in styles:
            style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = 'Courier New'
            style.font.size = Pt(10)
        if 'List Bullet' not in styles:
            style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.left_indent = Inches(0.25)
            style.paragraph_format.first_line_indent = Inches(-0.25)
        if 'List Number' not in styles:
            style = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
            style.paragraph_format.left_indent = Inches(0.25)
            style.paragraph_format.first_line_indent = Inches(-0.25)

    def apply_global_font(self):
        """将全局字体应用于所有默认样式，包括中文字体"""
        if not self.font_name_or_path:
            return
        styles = self.document.styles
        style_names = [
            'Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
            'List Bullet', 'List Number', 'Code'
        ]
        for style_name in style_names:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = self.font_name_or_path
                # 设置 XML 属性以确保中文字体生效
                if style._element.rPr.rFonts is None:
                    style._element.rPr.rFonts = OxmlElement('w:rFonts')
                style._element.rPr.rFonts.set(qn('w:ascii'), self.font_name_or_path)
                style._element.rPr.rFonts.set(qn('w:hAnsi'), self.font_name_or_path)
                style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name_or_path)
                style._element.rPr.rFonts.set(qn('w:cs'), self.font_name_or_path)

    def handle_element(self, element):
        if isinstance(element, NavigableString):
            self.handle_text(element.strip())
        elif element.name in self.tag_handlers:
            self.tag_handlers[element.name](element)
        else:
            for child in element.children:
                self.handle_element(child)

    def handle_text(self, text):
        if text and self.current_paragraph:
            run = self.current_paragraph.add_run(text)
            self.set_run_font(run)

    def handle_heading(self, element, level):
        self.current_paragraph = self.document.add_heading(level=level)
        self.apply_style(self.current_paragraph, element.get('style'))
        for child in element.children:
            self.handle_element(child)
        # 确保标题中的所有 run 使用全局字体
        for run in self.current_paragraph.runs:
            self.set_run_font(run)

    def handle_paragraph(self, element):
        self.current_paragraph = self.document.add_paragraph()
        self.apply_style(self.current_paragraph, element.get('style'))
        for child in element.children:
            self.handle_element(child)

    def handle_link(self, element):
        href = element.get('href')
        if href:
            if href.lower().endswith(('.jpg', '.png', '.gif')):
                try:
                    response = requests.get(href, timeout=5)
                    image_stream = BytesIO(response.content)
                    self.document.add_picture(image_stream, width=Inches(5))
                except Exception as e:
                    print(f"无法下载图片 {href}: {e}")
            else:
                if self.current_paragraph is None:
                    self.current_paragraph = self.document.add_paragraph()
                run = self.current_paragraph.add_run(element.get_text())
                run.font.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
                self.set_run_font(run)
                self.add_hyperlink(self.current_paragraph, href, element.get_text())

    def handle_strong(self, element):
        if self.current_paragraph:
            run = self.current_paragraph.add_run()
            run.bold = True
            self.set_run_font(run)
            for child in element.children:
                if isinstance(child, NavigableString):
                    run.add_text(child)
                else:
                    self.handle_element(child)

    def handle_em(self, element):
        if self.current_paragraph:
            run = self.current_paragraph.add_run()
            run.italic = True
            self.set_run_font(run)
            for child in element.children:
                if isinstance(child, NavigableString):
                    run.add_text(child)
                else:
                    self.handle_element(child)

    def handle_ul(self, element):
        self.list_type = 'bullet'
        self.list_level += 1
        for child in element.children:
            self.handle_element(child)
        self.list_level -= 1
        if self.list_level == 0:
            self.list_type = None

    def handle_ol(self, element):
        self.list_type = 'number'
        self.list_level += 1
        self.list_counters.append(0)
        for child in element.children:
            self.handle_element(child)
        self.list_counters.pop()
        self.list_level -= 1
        if self.list_level == 0:
            self.list_type = None

    def handle_li(self, element):
        if self.list_type == 'bullet':
            self.current_paragraph = self.document.add_paragraph(style='List Bullet')
        elif self.list_type == 'number':
            self.list_counters[-1] += 1
            self.current_paragraph = self.document.add_paragraph(style='List Number')
        else:
            self.current_paragraph = self.document.add_paragraph()
        indent = Inches(0.25 * self.list_level)
        self.current_paragraph.paragraph_format.left_indent = indent
        for child in element.children:
            self.handle_element(child)

    def handle_span(self, element):
        if self.current_paragraph:
            run = self.current_paragraph.add_run()
            self.apply_style(run, element.get('style'))
            self.set_run_font(run)
            for child in element.children:
                if isinstance(child, NavigableString):
                    run.add_text(child)
                else:
                    self.handle_element(child)

    def handle_br(self, element):
        if self.current_paragraph:
            self.current_paragraph.add_run('\n')

    def handle_hr(self, element):
        self.document.add_paragraph().add_run().add_break()
        self.current_paragraph = None

    def handle_header(self, element):
        section = self.document.sections[0]
        header = section.header
        self.current_paragraph = header.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_footer(self, element):
        section = self.document.sections[0]
        footer = section.footer
        self.current_paragraph = footer.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_nav(self, element):
        self.current_paragraph = self.document.add_paragraph()
        self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in element.children:
            self.handle_element(child)

    def handle_article(self, element):
        self.current_paragraph = self.document.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_section(self, element):
        self.current_paragraph = self.document.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_aside(self, element):
        self.current_paragraph = self.document.add_paragraph()
        self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for child in element.children:
            self.handle_element(child)

    def handle_figure(self, element):
        img = element.find('img')
        if img and img.get('src'):
            self.handle_img(img)

    def handle_figcaption(self, element):
        self.current_paragraph = self.document.add_paragraph()
        self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for child in element.children:
            self.handle_element(child)

    def handle_main(self, element):
        self.current_paragraph = self.document.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_mark(self, element):
        if self.current_paragraph:
            run = self.current_paragraph.add_run()
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            self.set_run_font(run)
            for child in element.children:
                if isinstance(child, NavigableString):
                    run.add_text(child)
                else:
                    self.handle_element(child)

    def handle_time(self, element):
        if self.current_paragraph:
            run = self.current_paragraph.add_run(element.get_text())
            self.set_run_font(run)

    def handle_pre(self, element):
        self.current_paragraph = self.document.add_paragraph(style='Code')
        for child in element.children:
            if isinstance(child, NavigableString):
                self.current_paragraph.add_run(child)
            else:
                self.handle_element(child)

    def handle_img(self, element):
        src = element.get('src')
        if src:
            try:
                response = requests.get(src, timeout=5)
                image_stream = BytesIO(response.content)
                self.document.add_picture(image_stream, width=Inches(5))
            except Exception as e:
                print(f"无法下载图片 {src}: {e}")

    def handle_table(self, element):
        rows = element.find_all('tr')
        if rows:
            table = self.document.add_table(rows=len(rows), cols=0)
            for row_idx, tr in enumerate(rows):
                self.handle_tr(tr, table, row_idx)

    def handle_tr(self, element, table, row_idx):
        cells = element.find_all(['td', 'th'])
        if cells and len(table.columns) < len(cells):
            for _ in range(len(cells) - len(table.columns)):
                table.add_column(Inches(1))
        for cell_idx, td in enumerate(cells):
            self.handle_td(td, table.cell(row_idx, cell_idx))

    def handle_td(self, element, cell):
        self.current_paragraph = cell.add_paragraph()
        for child in element.children:
            self.handle_element(child)

    def handle_th(self, element, cell):
        self.current_paragraph = cell.add_paragraph()
        run = self.current_paragraph.add_run(element.get_text())
        run.bold = True
        self.set_run_font(run)
        for child in element.children:
            self.handle_element(child)

    def set_run_font(self, run):
        """为 run 设置字体，确保中文字体生效"""
        run.font.name = self.font_name_or_path
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), self.font_name_or_path)
        rFonts.set(qn('w:hAnsi'), self.font_name_or_path)
        rFonts.set(qn('w:eastAsia'), self.font_name_or_path)
        rFonts.set(qn('w:cs'), self.font_name_or_path)

    def apply_style(self, docx_element, style_str):
        if not style_str:
            return
        sheet = cssutils.parseStyle(style_str)
        for prop in sheet:
            if prop.name == 'font-size':
                size = self.parse_font_size(prop.value)
                if hasattr(docx_element, 'font'):
                    docx_element.font.size = Pt(size)
                else:
                    docx_element.style.font.size = Pt(size)
            elif prop.name == 'color':
                color = self.parse_color(prop.value)
                if hasattr(docx_element, 'font'):
                    docx_element.font.color.rgb = color
            elif prop.name == 'font-weight' and prop.value == 'bold':
                if hasattr(docx_element, 'font'):
                    docx_element.font.bold = True
            elif prop.name == 'font-style' and prop.value == 'italic':
                if hasattr(docx_element, 'font'):
                    docx_element.font.italic = True

    def parse_font_size(self, value):
        if value.endswith('pt'):
            return float(value[:-2])
        elif value.endswith('px'):
            return float(value[:-2]) * 0.75
        return 12

    def parse_color(self, value):
        if value.startswith('#'):
            hex_color = value.lstrip('#')
            return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
        return RGBColor(0, 0, 0)

    def add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)